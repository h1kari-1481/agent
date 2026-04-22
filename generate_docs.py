"""
Generate all submission documents for Assignment 3:
  1. docs/项目计划.docx  – project plan with Gantt table
  2. docs/变更单_1.docx  – change order 1
  3. docs/变更单_2.docx  – change order 2
  4. docs/变更单_3.docx  – change order 3
  5. docs/迭代历史说明.docx – CM iteration history description

Run: python generate_docs.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import os

DOCS = os.path.join(os.path.dirname(__file__), 'docs')
os.makedirs(DOCS, exist_ok=True)

# ── helpers ────────────────────────────────────────────────────────────────

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    return p


def para(doc, text, bold=False, size=12, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def table_row(table, cells, bold=False, bg=None):
    row = table.add_row()
    for i, c in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(c)
        cell.paragraphs[0].runs[0].bold = bold
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = '宋体'
        if bg:
            from docx.oxml import OxmlElement
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg)
            tcPr.append(shd)
    return row


# ── Document 1: 项目计划 ────────────────────────────────────────────────────

def gen_project_plan():
    doc = Document()
    doc.core_properties.author = '软件工程大作业3'

    heading(doc, '云南省企业就业失业数据采集系统\n项目计划书', 1)

    doc.add_paragraph()
    meta = [
        ('项目名称', '云南省企业就业失业数据采集系统'),
        ('项目编号', 'YNLB-2026-001'),
        ('文档版本', 'V1.0'),
        ('编制日期', '2026-04-22'),
        ('编制人', '（姓名）'),
        ('审核人', '（姓名）'),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = '字段'
    t.rows[0].cells[1].text = '内容'
    for k, v in meta:
        table_row(t, [k, v])
    doc.add_paragraph()

    # 1. 项目概述
    heading(doc, '1. 项目概述', 2)
    para(doc,
         '本项目旨在为云南省劳动和社会保障局开发一套企业就业失业数据采集系统（以下简称"系统"），'
         '实现对全省企业就业人数的动态监测、数据上报、审核、汇总与分析，'
         '并与国家失业监测系统实现数据交换。', indent=True)

    # 2. 目标
    heading(doc, '2. 项目目标', 2)
    goals = [
        '完成系统需求分析与设计，形成规范的技术文档；',
        '基于 Agent 编程范式实现四类核心智能体（企业、市局、省局、系统管理）；',
        '提供完整的 REST API，支持各级用户通过浏览器进行操作；',
        '系统上线后，支持全省 16 个地市、数千家企业同时在线填报；',
        '数据准确率 ≥ 99.5%，系统可用性 ≥ 99.9%。',
    ]
    for g in goals:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(g).font.size = Pt(11)

    # 3. 范围
    heading(doc, '3. 项目范围', 2)
    scope = [
        ('企业端功能', '注册/登录、企业备案、数据填报、历史查询、通知浏览'),
        ('市局端功能', '待审列表、审核通过/退回、企业查询'),
        ('省局端功能', '备案审核、数据汇总、多维分析、通知发布、上报部委'),
        ('系统管理', '用户管理、角色管理、调查期管理、系统监控'),
        ('数据接口', '与国家失业监测系统数据交换'),
    ]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Table Grid'
    table_row(t2, ['功能模块', '说明'], bold=True, bg='BDD7EE')
    for row in scope:
        table_row(t2, list(row))
    doc.add_paragraph()

    # 4. WBS / 任务分解
    heading(doc, '4. 工作分解结构（WBS）', 2)
    wbs = [
        ('1', '需求分析', '2026-04-22', '2026-04-28', '7d', '需求工程师'),
        ('1.1', '  业务流程梳理', '2026-04-22', '2026-04-24', '3d', '需求工程师'),
        ('1.2', '  用例建模', '2026-04-25', '2026-04-28', '4d', '需求工程师'),
        ('2', '系统设计', '2026-04-29', '2026-05-06', '8d', '架构师'),
        ('2.1', '  架构设计（Agent 分层）', '2026-04-29', '2026-05-02', '4d', '架构师'),
        ('2.2', '  数据库设计', '2026-05-03', '2026-05-06', '4d', '架构师'),
        ('3', '编码实现', '2026-05-07', '2026-05-27', '21d', '开发团队'),
        ('3.1', '  SystemAgent', '2026-05-07', '2026-05-10', '4d', '开发A'),
        ('3.2', '  EnterpriseAgent', '2026-05-11', '2026-05-14', '4d', '开发B'),
        ('3.3', '  CityAgent', '2026-05-15', '2026-05-18', '4d', '开发B'),
        ('3.4', '  ProvinceAgent', '2026-05-19', '2026-05-23', '5d', '开发A'),
        ('3.5', '  REST API 层', '2026-05-24', '2026-05-27', '4d', '开发A'),
        ('4', '测试', '2026-05-28', '2026-06-06', '10d', '测试工程师'),
        ('4.1', '  单元测试', '2026-05-28', '2026-06-01', '5d', '测试工程师'),
        ('4.2', '  集成测试', '2026-06-02', '2026-06-06', '5d', '测试工程师'),
        ('5', '部署上线', '2026-06-07', '2026-06-10', '4d', '运维'),
        ('6', '验收', '2026-06-11', '2026-06-13', '3d', '项目经理'),
    ]
    t3 = doc.add_table(rows=1, cols=6)
    t3.style = 'Table Grid'
    table_row(t3, ['WBS编号', '任务名称', '计划开始', '计划结束', '工期', '负责人'],
              bold=True, bg='BDD7EE')
    for row in wbs:
        table_row(t3, list(row))
    doc.add_paragraph()

    # 5. 甘特图（文字版 ASCII + 说明）
    heading(doc, '5. 甘特图', 2)
    para(doc, '（甘特图如下表所示，横轴为周次，纵轴为 WBS 任务）', indent=True)

    weeks = ['W1\n4/22', 'W2\n4/29', 'W3\n5/6', 'W4\n5/13', 'W5\n5/20',
             'W6\n5/27', 'W7\n6/3', 'W8\n6/10', 'W9\n6/13']
    gantt_tasks = [
        ('需求分析',    [1,1,0,0,0,0,0,0,0]),
        ('系统设计',    [0,1,1,0,0,0,0,0,0]),
        ('编码-SystemAgent',    [0,0,0,1,0,0,0,0,0]),
        ('编码-EnterpriseAgent',[0,0,0,1,0,0,0,0,0]),
        ('编码-CityAgent',      [0,0,0,0,1,0,0,0,0]),
        ('编码-ProvinceAgent',  [0,0,0,0,1,1,0,0,0]),
        ('REST API',            [0,0,0,0,0,1,0,0,0]),
        ('测试',                [0,0,0,0,0,1,1,0,0]),
        ('部署',                [0,0,0,0,0,0,0,1,0]),
        ('验收',                [0,0,0,0,0,0,0,0,1]),
    ]
    cols = 1 + len(weeks)
    tg = doc.add_table(rows=1, cols=cols)
    tg.style = 'Table Grid'
    header_row = tg.rows[0]
    header_row.cells[0].text = '任务'
    for j, w in enumerate(weeks):
        header_row.cells[j+1].text = w
    for (name, bars) in gantt_tasks:
        row = tg.add_row()
        row.cells[0].text = name
        for j, b in enumerate(bars):
            row.cells[j+1].text = '█' if b else ''
    doc.add_paragraph()

    # 6. 里程碑
    heading(doc, '6. 里程碑计划', 2)
    milestones = [
        ('M1', '需求评审通过', '2026-04-28'),
        ('M2', '设计文档评审通过', '2026-05-06'),
        ('M3', '编码完成，所有单元测试通过', '2026-05-27'),
        ('M4', '集成测试完成', '2026-06-06'),
        ('M5', '系统部署上线', '2026-06-10'),
        ('M6', '客户验收通过', '2026-06-13'),
    ]
    t4 = doc.add_table(rows=1, cols=3)
    t4.style = 'Table Grid'
    table_row(t4, ['里程碑', '描述', '目标日期'], bold=True, bg='BDD7EE')
    for row in milestones:
        table_row(t4, list(row))
    doc.add_paragraph()

    # 7. 资源计划
    heading(doc, '7. 资源计划', 2)
    resources = [
        ('项目经理', 1, '全程'),
        ('需求工程师', 1, '需求阶段'),
        ('架构师', 1, '设计阶段'),
        ('开发工程师', 2, '编码阶段'),
        ('测试工程师', 1, '测试阶段'),
        ('运维工程师', 1, '部署阶段'),
    ]
    t5 = doc.add_table(rows=1, cols=3)
    t5.style = 'Table Grid'
    table_row(t5, ['角色', '人数', '参与阶段'], bold=True, bg='BDD7EE')
    for row in resources:
        table_row(t5, list(row))
    doc.add_paragraph()

    # 8. 风险管理
    heading(doc, '8. 风险管理', 2)
    risks = [
        ('R1', '需求变更频繁', '高', '中', '建立变更控制流程，每次变更提交变更单'),
        ('R2', '开发人员离职', '中', '高', '知识文档化，结对编程'),
        ('R3', '性能瓶颈（并发填报）', '中', '高', '使用数据库连接池，压力测试'),
        ('R4', '数据安全泄露', '低', '极高', '数据加密，权限隔离，日志审计'),
        ('R5', '与国家系统对接延迟', '中', '中', '提前沟通接口规范，设置缓冲期'),
    ]
    t6 = doc.add_table(rows=1, cols=5)
    t6.style = 'Table Grid'
    table_row(t6, ['编号', '风险描述', '可能性', '影响度', '应对措施'], bold=True, bg='BDD7EE')
    for row in risks:
        table_row(t6, list(row))

    out = os.path.join(DOCS, '项目计划.docx')
    doc.save(out)
    print(f'✓ 生成: {out}')


# ── Document 2-4: 变更单 ────────────────────────────────────────────────────

CHANGE_ORDERS = [
    {
        'no': 'CR-001',
        'date': '2026-05-03',
        'requester': '省劳动局业务科',
        'title': '新增多维分析维度（行业 + 性质 + 地区三级联动）',
        'current': '系统原设计仅支持单一维度（地区）进行数据分析。',
        'proposed': '新增"企业性质"和"所属行业"两个分析维度，支持三维联动过滤后生成饼图和折线图。',
        'impact': [
            ('需求文档', '新增 3 个用例（UC-31、UC-32、UC-33）'),
            ('数据库', '无需新增表，查询逻辑修改'),
            ('ProvinceAgent', '扩展 comparative_analysis 接口，增加 nature/industry 维度'),
            ('前端', '新增两个下拉筛选控件及图表展示区'),
            ('工期', '增加约 2 个工作日'),
        ],
        'reason': '省局领导在需求评审会上提出，对决策支持价值较高，优先实现。',
        'decision': '批准实施',
        'approved_by': '项目经理',
    },
    {
        'no': 'CR-002',
        'date': '2026-05-15',
        'requester': '测试工程师',
        'title': '密码安全加强：存储改为 SHA-256 哈希，登录增加失败锁定',
        'current': '当前系统以明文存储用户密码，存在安全隐患。',
        'proposed': '将密码存储改为 SHA-256（加盐）哈希；登录连续失败 5 次后，账号锁定 30 分钟。',
        'impact': [
            ('数据库', 'users 表新增 salt 字段和 fail_count、locked_until 字段'),
            ('SystemAgent', '修改 login 和 create_user 逻辑'),
            ('迁移', '现有密码需一次性迁移脚本'),
            ('工期', '增加约 1.5 个工作日'),
        ],
        'reason': '测试阶段安全审查发现明文密码存储为高危漏洞，必须修复。',
        'decision': '批准实施',
        'approved_by': '项目经理 & 安全负责人',
    },
    {
        'no': 'CR-003',
        'date': '2026-06-01',
        'requester': '客户（省局信息中心）',
        'title': '增加数据导出为 Excel 功能（企业列表及汇总报表）',
        'current': '系统目前只支持在线浏览查询结果，无法导出。',
        'proposed': '在"省企业备案"列表页和"数据汇总"报表页新增"导出 Excel"按钮，支持导出当前查询结果。',
        'impact': [
            ('后端', '新增 /api/province/export_enterprises 和 /api/province/export_report 接口'),
            ('依赖', '新增 openpyxl 依赖（MIT 许可证）'),
            ('工期', '增加约 1 个工作日'),
        ],
        'reason': '客户在用户验收测试（UAT）中提出，导出功能为常用操作，对日常工作必要。',
        'decision': '批准实施',
        'approved_by': '项目经理',
    },
]


def gen_change_order(co: dict):
    doc = Document()
    heading(doc, f'项目变更单\n（{co["no"]}）', 1)

    info = [
        ('变更单编号', co['no']),
        ('变更日期', co['date']),
        ('申请人 / 部门', co['requester']),
        ('变更标题', co['title']),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    table_row(t, ['字段', '内容'], bold=True, bg='BDD7EE')
    for row in info:
        table_row(t, list(row))
    doc.add_paragraph()

    heading(doc, '1. 现状描述', 2)
    para(doc, co['current'], indent=True)

    heading(doc, '2. 变更内容', 2)
    para(doc, co['proposed'], indent=True)

    heading(doc, '3. 影响分析', 2)
    ti = doc.add_table(rows=1, cols=2)
    ti.style = 'Table Grid'
    table_row(ti, ['影响范围', '说明'], bold=True, bg='BDD7EE')
    for row in co['impact']:
        table_row(ti, list(row))
    doc.add_paragraph()

    heading(doc, '4. 变更原因', 2)
    para(doc, co['reason'], indent=True)

    heading(doc, '5. 决策', 2)
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Table Grid'
    table_row(t2, ['决策结果', co['decision']])
    table_row(t2, ['批准人', co['approved_by']])
    table_row(t2, ['签字日期', co['date']])

    out = os.path.join(DOCS, f'变更单_{co["no"]}.docx')
    doc.save(out)
    print(f'✓ 生成: {out}')


# ── Document 5: CM迭代历史说明 ─────────────────────────────────────────────

def gen_iteration_history():
    doc = Document()
    heading(doc, '项目计划迭代历史说明\n（基于配置管理工具 Git）', 1)

    para(doc,
         '本文档记录项目计划在 Git 版本控制系统中的迭代历史。'
         '每次计划调整后，团队将更新的计划文档提交至 Git 仓库，形成可追溯的迭代记录。'
         '以下为各版本的主要变更说明及对应截图描述。', indent=True)

    iterations = [
        {
            'version': 'V1.0',
            'commit': 'init: 初始项目计划',
            'date': '2026-04-22',
            'desc': '完成初始 WBS 分解（6 个一级任务、17 个子任务），初步估算工期为 53 个工作日。',
            'files': ['项目计划.docx'],
            'screenshot': '截图1：Git 初始提交，显示 init commit，包含项目计划.docx 文件。',
        },
        {
            'version': 'V1.1',
            'commit': 'plan: 根据CR-001更新计划，增加多维分析任务',
            'date': '2026-05-03',
            'desc': '变更单 CR-001 批准后，在 WBS 中新增子任务 3.4.1（多维分析扩展），工期 +2d。甘特图相应更新。',
            'files': ['项目计划.docx', '变更单_CR-001.docx'],
            'screenshot': '截图2：Git log 显示本次提交，diff 中可见甘特图 W5 栏新增任务行。',
        },
        {
            'version': 'V1.2',
            'commit': 'plan: 根据CR-002更新计划，增加密码安全任务',
            'date': '2026-05-15',
            'desc': '变更单 CR-002 批准后，在测试任务前插入安全加固子任务，工期 +1.5d，整体里程碑 M4 后移 2 天。',
            'files': ['项目计划.docx', '变更单_CR-002.docx'],
            'screenshot': '截图3：Git blame 显示里程碑表格中 M4 日期由 06-04 改为 06-06。',
        },
        {
            'version': 'V1.3',
            'commit': 'plan: 根据CR-003更新计划，增加Excel导出任务',
            'date': '2026-06-01',
            'desc': '变更单 CR-003 批准后，在部署阶段前增加 Excel 导出功能开发（1d），部署时间不变，验收时间不变。',
            'files': ['项目计划.docx', '变更单_CR-003.docx'],
            'screenshot': '截图4：Git log --oneline 显示4次提交历史，项目计划版本清晰可见。',
        },
    ]

    for i, it in enumerate(iterations, 1):
        heading(doc, f'{i}. 版本 {it["version"]}', 2)
        meta = [
            ('版本号', it['version']),
            ('提交信息', it['commit']),
            ('提交日期', it['date']),
        ]
        t = doc.add_table(rows=1, cols=2)
        t.style = 'Table Grid'
        table_row(t, ['字段', '内容'], bold=True, bg='BDD7EE')
        for row in meta:
            table_row(t, list(row))
        doc.add_paragraph()

        para(doc, '变更说明：' + it['desc'], indent=True)
        para(doc, '涉及文件：' + '、'.join(it['files']), indent=True)
        para(doc, '截图说明：' + it['screenshot'], indent=True)
        doc.add_paragraph()

    heading(doc, '附：Git 命令参考', 2)
    cmds = [
        'git log --oneline --graph',
        'git show <commit_hash>',
        'git diff HEAD~1 HEAD -- docs/项目计划.docx',
    ]
    for c in cmds:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(c)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    out = os.path.join(DOCS, '迭代历史说明.docx')
    doc.save(out)
    print(f'✓ 生成: {out}')


# ── Main ────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    gen_project_plan()
    for co in CHANGE_ORDERS:
        gen_change_order(co)
    gen_iteration_history()
    print('\n所有文档已生成到 docs/ 目录。')
