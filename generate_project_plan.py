#!/usr/bin/env python3
"""
Generate 项目计划.docx for 云南省企业就业失业数据采集系统
Follows the 6-chapter structure of the reference document 软件项目管理案例.docx
"""

import os
import sys
import math
import datetime
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.font_manager as fm
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Font setup ───────────────────────────────────────────────────────────────
FONT_PATH_SC = None
for candidate in [
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc',
    '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
    '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
]:
    if os.path.exists(candidate):
        FONT_PATH_SC = candidate
        break

if FONT_PATH_SC:
    fm.fontManager.addfont(FONT_PATH_SC)

matplotlib.rcParams['font.sans-serif'] = ['Noto Sans CJK SC', 'Noto Sans CJK JP',
                                            'WenQuanYi Zen Hei', 'WenQuanYi Micro Hei',
                                            'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False

# ─── Output directory ─────────────────────────────────────────────────────────
IMG_DIR = '/home/runner/work/agent/agent/project_images'
os.makedirs(IMG_DIR, exist_ok=True)

def img(name):
    return os.path.join(IMG_DIR, name)

# ─── Date helpers ─────────────────────────────────────────────────────────────
def d(s):
    return datetime.datetime.strptime(s, '%Y-%m-%d')

def days(start, end):
    return (end - start).days

PROJECT_START = d('2026-04-22')
PROJECT_END   = d('2026-06-13')

# ─── WBS task data ────────────────────────────────────────────────────────────
TASKS = [
    # (id, name, start, end, level, phase)
    ('1',   '需求分析',              '2026-04-22','2026-04-28', 1, 'req'),
    ('1.1', '  业务流程梳理',         '2026-04-22','2026-04-24', 2, 'req'),
    ('1.2', '  用例建模',             '2026-04-25','2026-04-28', 2, 'req'),
    ('2',   '系统设计',              '2026-04-29','2026-05-06', 1, 'des'),
    ('2.1', '  架构设计（Agent分层）', '2026-04-29','2026-05-02', 2, 'des'),
    ('2.2', '  数据库设计',           '2026-05-03','2026-05-06', 2, 'des'),
    ('3',   '编码实现',              '2026-05-07','2026-05-27', 1, 'cod'),
    ('3.1', '  SystemAgent',         '2026-05-07','2026-05-10', 2, 'cod'),
    ('3.2', '  EnterpriseAgent',     '2026-05-11','2026-05-14', 2, 'cod'),
    ('3.3', '  CityAgent',           '2026-05-15','2026-05-18', 2, 'cod'),
    ('3.4', '  ProvinceAgent',       '2026-05-19','2026-05-23', 2, 'cod'),
    ('3.5', '  REST API层',           '2026-05-24','2026-05-27', 2, 'cod'),
    ('4',   '测试',                  '2026-05-28','2026-06-06', 1, 'tst'),
    ('4.1', '  单元测试',             '2026-05-28','2026-06-01', 2, 'tst'),
    ('4.2', '  集成测试',             '2026-06-02','2026-06-06', 2, 'tst'),
    ('5',   '部署上线',              '2026-06-07','2026-06-10', 1, 'dep'),
    ('6',   '验收',                  '2026-06-11','2026-06-13', 1, 'acc'),
]

PHASE_COLORS = {
    'req': '#4472C4',
    'des': '#70AD47',
    'cod': '#FFC000',
    'tst': '#FF0000',
    'dep': '#7030A0',
    'acc': '#808080',
}

MILESTONES = [
    ('M1', '需求评审通过',       '2026-04-28'),
    ('M2', '设计文档评审通过',    '2026-05-06'),
    ('M3', '编码完成/单测通过',   '2026-05-27'),
    ('M4', '集成测试完成',        '2026-06-06'),
    ('M5', '系统部署上线',        '2026-06-10'),
    ('M6', '客户验收通过',        '2026-06-13'),
]

# ═══════════════════════════════════════════════════════════════════════════════
# 1. GANTT CHART
# ═══════════════════════════════════════════════════════════════════════════════
def make_gantt():
    fig, ax = plt.subplots(figsize=(18, 11))

    task_names = [t[1] for t in TASKS]
    n = len(TASKS)

    ref = PROJECT_START
    for i, (tid, tname, ts, te, lvl, phase) in enumerate(TASKS):
        s = days(ref, d(ts))
        dur = days(d(ts), d(te)) + 1
        color = PHASE_COLORS[phase]
        alpha = 1.0 if lvl == 1 else 0.65
        ax.broken_barh([(s, dur)], (n - i - 0.7, 0.5),
                       facecolors=color, edgecolors='white', alpha=alpha, linewidth=0.5)
        ax.text(s + dur + 0.3, n - i - 0.45, f'{dur}d',
                va='center', ha='left', fontsize=7.5, color='#333333')

    # Milestone diamonds
    milestone_rows = {
        '2026-04-28': 2,   # row for 需求分析
        '2026-05-06': 4,   # 系统设计
        '2026-05-27': 6,   # 编码实现
        '2026-06-06': 8,   # 测试
        '2026-06-10': 9,   # 部署
        '2026-06-13': 10,  # 验收
    }
    for mid, mdesc, mdate in MILESTONES:
        x = days(ref, d(mdate))
        row_i = next((i for i, t in enumerate(TASKS) if t[2] <= mdate <= t[3] and t[4] == 1), 0)
        y = n - row_i - 0.45
        ax.plot(x + 1, y, 'D', color='red', markersize=8, zorder=5)
        ax.text(x + 1.5, y + 0.15, f'{mid}', fontsize=7.5, color='red', fontweight='bold')

    # X-axis tick marks = dates
    total_days = days(ref, PROJECT_END) + 1
    tick_dates = []
    cur = ref
    while cur <= PROJECT_END:
        tick_dates.append(cur)
        cur += datetime.timedelta(days=7)
    tick_dates.append(PROJECT_END)

    ax.set_xticks([days(ref, td) for td in tick_dates])
    ax.set_xticklabels([td.strftime('%m/%d') for td in tick_dates], rotation=45, fontsize=8)
    ax.set_xlim(-1, total_days + 8)

    ax.set_yticks(range(n))
    ax.set_yticklabels([t[1] for t in reversed(TASKS)], fontsize=9)
    ax.set_ylim(-0.5, n)

    ax.set_xlabel('日期', fontsize=10)
    ax.set_title('项目甘特图\n云南省企业就业失业数据采集系统', fontsize=13, fontweight='bold', pad=12)

    # Grid
    ax.xaxis.grid(True, linestyle='--', alpha=0.4)
    ax.set_axisbelow(True)

    # Legend
    legend_items = [
        mpatches.Patch(color=PHASE_COLORS['req'], label='需求分析'),
        mpatches.Patch(color=PHASE_COLORS['des'], label='系统设计'),
        mpatches.Patch(color=PHASE_COLORS['cod'], label='编码实现'),
        mpatches.Patch(color=PHASE_COLORS['tst'], label='测试'),
        mpatches.Patch(color=PHASE_COLORS['dep'], label='部署上线'),
        mpatches.Patch(color=PHASE_COLORS['acc'], label='验收'),
        plt.Line2D([0],[0], marker='D', color='w', markerfacecolor='red',
                   markersize=8, label='里程碑'),
    ]
    ax.legend(handles=legend_items, loc='lower right', fontsize=8, ncol=4)

    plt.tight_layout()
    out = img('gantt.png')
    plt.savefig(out, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'✓ Gantt chart saved: {out}')
    return out

# ═══════════════════════════════════════════════════════════════════════════════
# 2. NETWORK DIAGRAM
# ═══════════════════════════════════════════════════════════════════════════════
def make_network():
    fig, ax = plt.subplots(figsize=(16, 9))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_facecolor('#F8F9FA')
    fig.patch.set_facecolor('#F8F9FA')
    ax.set_title('项目网络图（关键路径法）\n云南省企业就业失业数据采集系统',
                 fontsize=13, fontweight='bold', pad=12)

    # Node layout (x, y, id, label, duration, early_start, early_finish)
    nodes = [
        (1.5, 4.5, '开始',   '开始',      '',  '',   ''),
        (3.8, 7.0, '1.1',   '业务流程\n梳理',   '3d', 'ES:0',  'EF:3'),
        (3.8, 2.0, '1.2',   '用例建模',   '4d', 'ES:3',  'EF:7'),
        (6.5, 4.5, '2.1',   '架构设计\n(Agent)', '4d', 'ES:7',  'EF:11'),
        (6.5, 2.0, '2.2',   '数据库设计', '4d', 'ES:11', 'EF:15'),
        (9.2, 7.0, '3.1',   'SystemAgent','4d', 'ES:15', 'EF:19'),
        (9.2, 5.5, '3.2',   'Enterprise\nAgent','4d','ES:19','EF:23'),
        (9.2, 4.0, '3.3',   'CityAgent',  '4d', 'ES:23', 'EF:27'),
        (9.2, 2.5, '3.4',   'Province\nAgent','5d','ES:27','EF:32'),
        (9.2, 1.0, '3.5',   'REST API层', '4d', 'ES:32', 'EF:36'),
        (12.0,6.0, '4.1',   '单元测试',   '5d', 'ES:36', 'EF:41'),
        (12.0,3.5, '4.2',   '集成测试',   '5d', 'ES:41', 'EF:46'),
        (14.0,4.5, '5',     '部署上线',   '4d', 'ES:46', 'EF:50'),
        (15.5,4.5, '结束',   '结束',      '',   '',   ''),
    ]

    # Draw nodes
    node_pos = {}
    for (x, y, nid, label, dur, es, ef) in nodes:
        node_pos[nid] = (x, y)
        is_start_end = nid in ('开始', '结束')
        if is_start_end:
            circ = plt.Circle((x, y), 0.45, color='#2E75B6', zorder=3)
            ax.add_patch(circ)
            ax.text(x, y, label, ha='center', va='center',
                    fontsize=8, color='white', fontweight='bold', zorder=4)
        else:
            rect = mpatches.FancyBboxPatch((x - 0.75, y - 0.70), 1.5, 1.4,
                                           boxstyle='round,pad=0.05',
                                           facecolor='#E9F0FB', edgecolor='#2E75B6',
                                           linewidth=1.5, zorder=3)
            ax.add_patch(rect)
            # Top: task name
            ax.text(x, y + 0.30, label, ha='center', va='center',
                    fontsize=7.2, fontweight='bold', color='#1F3864', zorder=4)
            # Bottom: ES/EF
            ax.text(x - 0.35, y - 0.42, es, ha='center', va='center',
                    fontsize=6.5, color='#555555', zorder=4)
            ax.text(x + 0.35, y - 0.42, ef, ha='center', va='center',
                    fontsize=6.5, color='#555555', zorder=4)
            # Duration label
            ax.text(x, y - 0.05, dur, ha='center', va='center',
                    fontsize=7, color='#C00000', fontweight='bold', zorder=4)

    # Edges (from, to, is_critical)
    edges = [
        ('开始', '1.1', True),
        ('开始', '1.2', False),
        ('1.1', '2.1', True),
        ('1.2', '2.1', False),
        ('2.1', '2.2', True),
        ('2.2', '3.1', True),
        ('2.2', '3.2', True),
        ('2.2', '3.3', True),
        ('2.2', '3.4', True),
        ('2.2', '3.5', True),
        ('3.1', '4.1', True),
        ('3.2', '4.1', False),
        ('3.3', '4.1', False),
        ('3.4', '4.1', False),
        ('3.5', '4.1', False),
        ('4.1', '4.2', True),
        ('4.2', '5',   True),
        ('5',   '结束', True),
    ]

    for (fr, to, crit) in edges:
        x1, y1 = node_pos[fr]
        x2, y2 = node_pos[to]
        color = '#C00000' if crit else '#888888'
        lw    = 2.0 if crit else 1.0
        ax.annotate('', xy=(x2 - 0.75, y2), xytext=(x1 + 0.75, y1),
                    arrowprops=dict(arrowstyle='->', color=color, lw=lw),
                    zorder=2)

    # Critical path legend
    ax.plot([], [], '-', color='#C00000', lw=2, label='关键路径')
    ax.plot([], [], '-', color='#888888', lw=1, label='非关键路径')
    ax.legend(loc='lower right', fontsize=9)

    # Annotation
    ax.text(8, 0.3, '关键路径总工期：53天（2026-04-22 → 2026-06-13）',
            ha='center', fontsize=9, color='#C00000',
            bbox=dict(boxstyle='round', facecolor='#FFE6E6', edgecolor='#C00000'))

    plt.tight_layout()
    out = img('network.png')
    plt.savefig(out, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'✓ Network diagram saved: {out}')
    return out

# ═══════════════════════════════════════════════════════════════════════════════
# 3. MILESTONE CHART
# ═══════════════════════════════════════════════════════════════════════════════
def make_milestone():
    fig, ax = plt.subplots(figsize=(14, 5))

    ref = PROJECT_START
    xs = [days(ref, d(m[2])) for m in MILESTONES]
    labels = [f"{m[0]}\n{m[1]}\n{m[2]}" for m in MILESTONES]

    ax.axhline(y=0, color='#2E75B6', lw=2, zorder=1)

    # Alternating above/below labels
    for i, (x, lab) in enumerate(zip(xs, labels)):
        y_dir = 1 if i % 2 == 0 else -1
        y_lbl = 1.6 if y_dir == 1 else -1.6
        y_con = 0.15 * y_dir

        ax.annotate('', xy=(x, y_con), xytext=(x, y_lbl - 0.25 * y_dir),
                    arrowprops=dict(arrowstyle='-', color='#666666', lw=1))
        ax.plot(x, 0, 'D', color='#C00000', markersize=14, zorder=4)
        ax.text(x, 0, MILESTONES[i][0], ha='center', va='center',
                fontsize=8, color='white', fontweight='bold', zorder=5)
        ax.text(x, y_lbl, lab, ha='center', va='center' if y_dir == 1 else 'top',
                fontsize=8.5, color='#1F3864',
                bbox=dict(boxstyle='round,pad=0.3', facecolor='#DEEAF1',
                          edgecolor='#2E75B6', linewidth=1))

    total = days(ref, PROJECT_END)
    ax.set_xlim(-5, total + 5)
    ax.set_ylim(-3.2, 3.2)

    # X ticks every 2 weeks
    tick_days = list(range(0, total + 1, 14)) + [total]
    ax.set_xticks(tick_days)
    ax.set_xticklabels([(ref + datetime.timedelta(days=td)).strftime('%m/%d')
                        for td in tick_days], fontsize=8.5)
    ax.set_yticks([])
    ax.set_xlabel('日期', fontsize=10)
    ax.set_title('项目里程碑图\n云南省企业就业失业数据采集系统',
                 fontsize=13, fontweight='bold', pad=10)

    # Remove spines except bottom
    for spine in ['top', 'right', 'left']:
        ax.spines[spine].set_visible(False)
    ax.spines['bottom'].set_visible(False)

    plt.tight_layout()
    out = img('milestone.png')
    plt.savefig(out, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'✓ Milestone chart saved: {out}')
    return out

# ═══════════════════════════════════════════════════════════════════════════════
# 4. WBS DIAGRAM
# ═══════════════════════════════════════════════════════════════════════════════
def make_wbs():
    fig, ax = plt.subplots(figsize=(18, 9))
    ax.set_xlim(0, 18)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_facecolor('#FAFAFA')
    fig.patch.set_facecolor('#FAFAFA')
    ax.set_title('工作分解结构（WBS）\n云南省企业就业失业数据采集系统',
                 fontsize=13, fontweight='bold', pad=12)

    def box(ax, x, y, w, h, text, facecolor='#DEEAF1', edgecolor='#2E75B6',
            fontsize=8.5, bold=False):
        rect = mpatches.FancyBboxPatch((x - w/2, y - h/2), w, h,
                                       boxstyle='round,pad=0.05',
                                       facecolor=facecolor, edgecolor=edgecolor,
                                       linewidth=1.5, zorder=3)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
                fontweight='bold' if bold else 'normal', color='#1F3864', zorder=4,
                wrap=True)

    def line(ax, x1, y1, x2, y2):
        ax.plot([x1, x2], [y1, y2], '-', color='#2E75B6', lw=1.2, zorder=2)

    # Root
    root_x, root_y = 9, 8.2
    box(ax, root_x, root_y, 4.5, 0.8,
        '云南省企业就业失业\n数据采集系统',
        facecolor='#2E75B6', edgecolor='#1F3864', fontsize=9, bold=True)
    # Override text color for dark bg
    ax.texts[-1].set_color('white')

    # Level 1 phases
    lvl1 = [
        (2.0, 5.8, '1\n需求分析',  '#70AD47'),
        (5.2, 5.8, '2\n系统设计',  '#4472C4'),
        (8.0, 5.8, '3\n编码实现',  '#FFC000'),
        (11.0,5.8, '4\n测试',      '#FF4444'),
        (14.2,5.8, '5\n部署上线',  '#7030A0'),
        (16.8,5.8, '6\n验收',      '#808080'),
    ]
    for (x, y, text, color) in lvl1:
        line(ax, root_x, root_y - 0.4, x, y + 0.4)
        box(ax, x, y, 2.4, 0.75, text, facecolor=color, edgecolor='white', fontsize=8.5, bold=True)
        ax.texts[-1].set_color('white')

    # Level 2
    lvl2 = {
        '1': [(1.1, 3.9, '1.1\n业务流程梳理'), (3.0, 3.9, '1.2\n用例建模')],
        '2': [(4.3, 3.9, '2.1\n架构设计\n(Agent分层)'), (6.1, 3.9, '2.2\n数据库设计')],
        '3': [(7.0, 3.9, '3.1\nSystemAgent'), (8.0, 3.9, '3.2\nEnterprise\nAgent'),
              (9.0, 3.9, '3.3\nCityAgent'), (10.0,3.9, '3.4\nProvince\nAgent'),
              (11.0,3.9, '3.5\nREST API层')],
        '4': [(10.3,3.9, '4.1\n单元测试'), (11.8,3.9, '4.2\n集成测试')],
    }
    parent_x = {
        '1': 2.0, '2': 5.2, '3': 8.0, '4': 11.0,
    }
    colors2 = {
        '1': '#E2EFD9', '2': '#DEEAF1', '3': '#FFF2CC', '4': '#FFDEDE',
    }
    for pid, children in lvl2.items():
        px = parent_x[pid]
        for (cx, cy, ctext) in children:
            line(ax, px, 5.8 - 0.38, cx, cy + 0.45)
            box(ax, cx, cy, 1.7, 0.85, ctext,
                facecolor=colors2[pid], edgecolor='#2E75B6', fontsize=7.5)

    plt.tight_layout()
    out = img('wbs.png')
    plt.savefig(out, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'✓ WBS diagram saved: {out}')
    return out

# ═══════════════════════════════════════════════════════════════════════════════
# 5. LIFECYCLE MODEL (WATERFALL)
# ═══════════════════════════════════════════════════════════════════════════════
def make_lifecycle():
    fig, ax = plt.subplots(figsize=(14, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_facecolor('#F8F9FA')
    fig.patch.set_facecolor('#F8F9FA')
    ax.set_title('软件生命周期模型（瀑布模型）\n云南省企业就业失业数据采集系统',
                 fontsize=13, fontweight='bold', pad=12)

    phases = [
        ('软件规划', '#4472C4', '确定项目范围、可行性分析、\n资源与进度规划'),
        ('需求开发', '#70AD47', '业务流程梳理、用例建模、\n需求规格说明书'),
        ('软件结构设计', '#FFC000', '系统架构设计、Agent分层设计、\n接口设计'),
        ('数据库设计', '#ED7D31', '数据模型、ER图、\n表结构设计'),
        ('实施（编码）', '#FF4444', '四类Agent实现、REST API层、\n前端页面开发'),
        ('系统集成', '#7030A0', '模块集成、单元测试、\n集成测试'),
        ('提交上线', '#00B0F0', '部署到生产环境、\n用户培训、系统验收'),
        ('维护', '#808080', '缺陷修复、版本升级、\n运营支持'),
    ]

    n = len(phases)
    # Staircase waterfall
    for i, (name, color, desc) in enumerate(phases):
        x = i * 1.5 + 0.3
        y = 5.5 - i * 0.55
        w, h = 2.2, 0.7

        # Parallelogram-like effect using FancyBboxPatch
        rect = mpatches.FancyBboxPatch((x, y), w, h,
                                       boxstyle='round,pad=0.04',
                                       facecolor=color, edgecolor='white',
                                       linewidth=1.5, alpha=0.9, zorder=3)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, name, ha='center', va='center',
                fontsize=9, fontweight='bold', color='white', zorder=4)

        # Description text
        ax.text(x + w/2, y - 0.42, desc, ha='center', va='top',
                fontsize=7, color='#444444', zorder=4)

        # Arrow to next
        if i < n - 1:
            x2 = (i+1) * 1.5 + 0.3
            y2 = 5.5 - (i+1) * 0.55 + h
            ax.annotate('', xy=(x2 + w/2, y2),
                        xytext=(x + w/2, y),
                        arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5),
                        zorder=2)

    # Feedback arrow (maintenance back to requirement)
    ax.annotate('', xy=(0.3 + 2.2/2, 5.5 + 0.7),
                xytext=(n * 1.5 - 1.5 + 0.3 + 2.2/2, 5.5 - (n-1)*0.55),
                arrowprops=dict(arrowstyle='->', color='#C00000', lw=1.5,
                                connectionstyle='arc3,rad=0.3'),
                zorder=2)
    ax.text(7, 6.5, '反馈/变更', ha='center', fontsize=8, color='#C00000')

    plt.tight_layout()
    out = img('lifecycle.png')
    plt.savefig(out, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'✓ Lifecycle model saved: {out}')
    return out

# ═══════════════════════════════════════════════════════════════════════════════
# 6. ORG CHART
# ═══════════════════════════════════════════════════════════════════════════════
def make_orgchart():
    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_facecolor('#FAFAFA')
    fig.patch.set_facecolor('#FAFAFA')
    ax.set_title('项目组织结构图\n云南省企业就业失业数据采集系统',
                 fontsize=13, fontweight='bold', pad=12)

    def box(x, y, w, h, text, fc='#DEEAF1', ec='#2E75B6', fs=9, bold=False, tc='#1F3864'):
        rect = mpatches.FancyBboxPatch((x - w/2, y - h/2), w, h,
                                       boxstyle='round,pad=0.06',
                                       facecolor=fc, edgecolor=ec,
                                       linewidth=1.5, zorder=3)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=fs,
                fontweight='bold' if bold else 'normal', color=tc, zorder=4)

    def line(x1, y1, x2, y2):
        ax.plot([x1, x2], [y1, y2], '-', color='#2E75B6', lw=1.3, zorder=2)

    # Root
    box(6.5, 6.2, 3.0, 0.75, '项目经理', fc='#2E75B6', ec='#1F3864',
        fs=11, bold=True, tc='white')

    # Level 2
    lvl2 = [
        (2.0, 4.6, '需求工程师', '#70AD47'),
        (4.5, 4.6, '架构师',    '#FFC000'),
        (7.5, 4.6, '开发团队',  '#4472C4'),
        (10.5,4.6, '测试工程师','#FF4444'),
        (12.5,4.6, '运维工程师','#7030A0'),
    ]
    for (x, y, name, color) in lvl2:
        line(6.5, 6.2 - 0.38, x, y + 0.38)
        box(x, y, 2.2, 0.72, name, fc=color, ec='white', fs=9.5, bold=True, tc='white')

    # Dev sub-team
    devs = [(6.2, 3.1, '开发工程师A'), (8.8, 3.1, '开发工程师B')]
    for (x, y, name) in devs:
        line(7.5, 4.6 - 0.36, x, y + 0.36)
        box(x, y, 2.2, 0.65, name, fc='#DEEAF1', ec='#2E75B6', fs=9)

    # Responsibility labels
    resp = {
        '需求工程师': '业务流程梳理\n用例建模',
        '架构师':    '架构设计\n数据库设计',
        '开发工程师A': 'SystemAgent\nProvinceAgent\nREST API',
        '开发工程师B': 'EnterpriseAgent\nCityAgent',
        '测试工程师': '单元/集成测试',
        '运维工程师': '部署上线\n运维保障',
    }
    positions = {
        '需求工程师': (2.0, 3.4),
        '架构师':    (4.5, 3.4),
        '开发工程师A': (6.2, 2.2),
        '开发工程师B': (8.8, 2.2),
        '测试工程师': (10.5, 3.4),
        '运维工程师': (12.5, 3.4),
    }
    for name, (rx, ry) in positions.items():
        ax.text(rx, ry, resp[name], ha='center', va='top', fontsize=7.2,
                color='#555555', style='italic',
                bbox=dict(boxstyle='round,pad=0.2', facecolor='#F5F5F5',
                          edgecolor='#CCCCCC', linewidth=0.8))

    plt.tight_layout()
    out = img('orgchart.png')
    plt.savefig(out, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'✓ Org chart saved: {out}')
    return out

# ═══════════════════════════════════════════════════════════════════════════════
# 7. RESOURCE HISTOGRAM
# ═══════════════════════════════════════════════════════════════════════════════
def make_resource_histogram():
    # Periods (bi-weekly)
    periods = [
        ('4/22-4/28', 1, 0),    # (label, req, des, cod, tst, dep, acc)
        ('4/29-5/6',  0, 1),
        ('5/7-5/13',  0, 0),
        ('5/14-5/20', 0, 0),
        ('5/21-5/27', 0, 0),
        ('5/28-6/6',  0, 0),
        ('6/7-6/10',  0, 0),
        ('6/11-6/13', 0, 0),
    ]

    labels  = ['4/22-4/28','4/29-5/6','5/7-5/13','5/14-5/20',
               '5/21-5/27','5/28-6/6','6/7-6/10','6/11-6/13']

    # Rows: 项目经理, 需求工程师, 架构师, 开发A, 开发B, 测试工程师, 运维工程师
    # Columns = periods
    resource_data = {
        '项目经理':    [1, 1, 1, 1, 1, 1, 1, 1],
        '需求工程师':  [1, 0, 0, 0, 0, 0, 0, 0],
        '架构师':      [0, 1, 0, 0, 0, 0, 0, 0],
        '开发工程师A': [0, 0, 1, 1, 1, 0, 0, 0],
        '开发工程师B': [0, 0, 1, 1, 1, 0, 0, 0],
        '测试工程师':  [0, 0, 0, 0, 0, 1, 0, 0],
        '运维工程师':  [0, 0, 0, 0, 0, 0, 1, 1],
    }

    roles  = list(resource_data.keys())
    colors = ['#2E75B6','#70AD47','#FFC000','#ED7D31','#FF4444','#7030A0','#00B0F0']

    x = np.arange(len(labels))
    width = 0.10
    n = len(roles)

    fig, ax = plt.subplots(figsize=(14, 6))
    for i, (role, color) in enumerate(zip(roles, colors)):
        vals = resource_data[role]
        offset = (i - n/2 + 0.5) * width
        bars = ax.bar(x + offset, vals, width * 0.9, label=role,
                      color=color, alpha=0.85)

    # Total people line
    totals = [sum(resource_data[r][j] for r in roles) for j in range(len(labels))]
    ax.plot(x, totals, 'k--o', linewidth=1.5, markersize=6, label='合计人数', zorder=5)
    for i, t in enumerate(totals):
        ax.text(i, t + 0.08, str(t), ha='center', fontsize=9, fontweight='bold')

    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.set_yticks(range(0, max(totals) + 2))
    ax.set_ylabel('人数', fontsize=10)
    ax.set_xlabel('时间段', fontsize=10)
    ax.set_title('人力资源直方图\n云南省企业就业失业数据采集系统',
                 fontsize=13, fontweight='bold', pad=10)
    ax.legend(loc='upper right', fontsize=8, ncol=4)
    ax.yaxis.grid(True, linestyle='--', alpha=0.5)
    ax.set_axisbelow(True)

    plt.tight_layout()
    out = img('resource_histogram.png')
    plt.savefig(out, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'✓ Resource histogram saved: {out}')
    return out

# ═══════════════════════════════════════════════════════════════════════════════
# HELPER: DOCX formatting utilities
# ═══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trPr.append(trHeight)

def bold_cell(cell, text, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
              color='000000', bg=None):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bg:
        set_cell_bg(cell, bg)
    return run

def normal_cell(cell, text, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    h.paragraph_format.space_after  = Pt(6)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(doc, text, size=11, before=0, after=4, indent=0, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_image(doc, path, width_inches=5.8, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        for run in cp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# BUILD THE DOCX
# ═══════════════════════════════════════════════════════════════════════════════
def build_docx(images):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Cover Page ─────────────────────────────────────────────────────────────
    # Title paragraph
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(48)
    t.paragraph_format.space_after  = Pt(10)
    run = t.add_run('云南省企业就业失业数据采集系统')
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.space_after = Pt(36)
    run2 = t2.add_run('软件项目计划书')
    run2.bold = True
    run2.font.size = Pt(18)
    run2.font.name = '黑体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # Cover table
    tbl = doc.add_table(rows=7, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_data = [
        ('项目名称', '云南省企业就业失业数据采集系统', '项目编号', 'YNEM-2026-001'),
        ('委托单位', '云南省劳动和社会保障局',         '版本号',   'V1.0'),
        ('项目经理', '（待定）',                       '密级',     '内部'),
        ('编制日期', '2026年4月22日',                  '计划周期', '2026-04-22 至 2026-06-13'),
        ('编制人员', '项目管理团队',                   '批准人',   '（待定）'),
        ('文档状态', '正式发布',                       '页数',     '共若干页'),
    ]
    HEADER_BG = '2E75B6'
    for r_idx, row_data in enumerate(cover_data):
        row = tbl.rows[r_idx]
        set_row_height(row, 0.9)
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            if c_idx % 2 == 0:
                bold_cell(cell, text, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                          color='FFFFFF', bg=HEADER_BG)
            else:
                normal_cell(cell, text, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    # Last row - approval signature row
    last_row = tbl.rows[6]
    set_row_height(last_row, 1.5)
    last_row.cells[0].merge(last_row.cells[3])
    bold_cell(last_row.cells[0], '批准签字：_______________          日期：_______________',
              font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color='000000', bg='F2F2F2')

    doc.add_page_break()

    # ── Chapter 1: 前言 ────────────────────────────────────────────────────────
    add_heading(doc, '第一章  前言', 1)

    add_heading(doc, '1.1 项目开发背景', 2)
    add_para(doc,
        '随着云南省经济社会的持续发展，劳动力市场管理日趋复杂，就业与失业数据的实时监测与上报已成为'
        '劳动保障工作的重要环节。目前，省内各企业的就业失业数据依赖人工填报，存在数据滞后、准确性低、'
        '汇总效率低下等突出问题，难以满足云南省劳动和社会保障局对劳动力市场动态监测的需求。')
    add_para(doc,
        '国家人力资源和社会保障部要求各省建立与国家失业监测系统的数据交换机制，云南省亟需建立一套'
        '覆盖省、市、企业三级的就业失业数据采集、审核、汇总与分析系统，以实现数据的动态更新和'
        '与国家系统的互联互通。')
    add_para(doc,
        '基于上述背景，云南省劳动和社会保障局委托开发"云南省企业就业失业数据采集系统"，采用先进的'
        'Agent编程范式，构建企业、市、省、系统四类智能代理协同工作的软件架构，以满足多级用户的'
        '数据采集与管理需求。')

    add_heading(doc, '1.2 项目开发目的', 2)
    add_para(doc, '本项目的开发目的包括以下几个方面：')
    for item in [
        '（1）实现企业就业人数的动态监测：支持企业用户实时录入、修改和上报就业失业数据，确保数据的实时性与准确性。',
        '（2）建立三级数据管理体系：支持企业用户、市级用户和省级用户分级管理数据，实现从企业到市、再到省的逐级审核与汇总。',
        '（3）实现数据交换与共享：与国家失业监测系统对接，按照规定的接口规范实现数据的自动上传和同步。',
        '（4）提供数据分析与决策支持：通过多维度数据统计分析，为劳动保障部门的宏观决策提供科学依据。',
        '（5）提升管理效率：以信息化手段替代传统人工填报模式，大幅降低人力成本，提高工作效率。',
    ]:
        add_para(doc, item, indent=0.5)

    add_heading(doc, '1.3 项目开发意义', 2)
    add_para(doc,
        '本项目的成功实施将对云南省劳动力市场管理产生深远影响：')
    for item in [
        '社会意义：推动云南省就业保障信息化建设，提升政府公共服务能力，为广大劳动者提供更及时、便捷的就业服务。',
        '经济意义：减少人工数据处理成本，提高政策制定效率，助力劳动力资源的优化配置。',
        '技术意义：率先在省级劳动保障系统引入Agent编程范式，探索智能化政务系统建设的新路径，为全国同类系统提供参考。',
        '管理意义：建立标准化的数据采集和审核流程，强化数据治理，提升省内就业失业数据的可信度和权威性。',
    ]:
        add_para(doc, f'• {item}', indent=0.5)

    doc.add_page_break()

    # ── Chapter 2: 范围计划 ───────────────────────────────────────────────────
    add_heading(doc, '第二章  范围计划', 1)

    add_heading(doc, '2.1 项目工作分解结构（WBS）', 2)
    add_para(doc,
        '本项目按照软件工程标准进行工作分解，将项目总体工作分解为6个一级任务包和若干二级任务包。'
        '下图为项目工作分解结构（WBS）图示。')
    add_image(doc, images['wbs'], width_inches=6.2,
              caption='图2-1  项目工作分解结构（WBS）')

    add_para(doc, '各WBS任务包说明如下：')

    # WBS table
    wbs_tbl = doc.add_table(rows=1, cols=6)
    wbs_tbl.style = 'Table Grid'
    wbs_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['WBS编号', '任务名称', '计划开始', '计划结束', '工期', '负责人']
    for i, h in enumerate(headers):
        bold_cell(wbs_tbl.rows[0].cells[i], h, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')

    wbs_rows = [
        ('1',   '需求分析',              '2026-04-22','2026-04-28','7天','需求工程师'),
        ('1.1', '  业务流程梳理',         '2026-04-22','2026-04-24','3天','需求工程师'),
        ('1.2', '  用例建模',             '2026-04-25','2026-04-28','4天','需求工程师'),
        ('2',   '系统设计',              '2026-04-29','2026-05-06','8天','架构师'),
        ('2.1', '  架构设计（Agent分层）','2026-04-29','2026-05-02','4天','架构师'),
        ('2.2', '  数据库设计',           '2026-05-03','2026-05-06','4天','架构师'),
        ('3',   '编码实现',              '2026-05-07','2026-05-27','21天','开发团队'),
        ('3.1', '  SystemAgent',         '2026-05-07','2026-05-10','4天','开发A'),
        ('3.2', '  EnterpriseAgent',     '2026-05-11','2026-05-14','4天','开发B'),
        ('3.3', '  CityAgent',           '2026-05-15','2026-05-18','4天','开发B'),
        ('3.4', '  ProvinceAgent',       '2026-05-19','2026-05-23','5天','开发A'),
        ('3.5', '  REST API层',           '2026-05-24','2026-05-27','4天','开发A'),
        ('4',   '测试',                  '2026-05-28','2026-06-06','10天','测试工程师'),
        ('4.1', '  单元测试',             '2026-05-28','2026-06-01','5天','测试工程师'),
        ('4.2', '  集成测试',             '2026-06-02','2026-06-06','5天','测试工程师'),
        ('5',   '部署上线',              '2026-06-07','2026-06-10','4天','运维'),
        ('6',   '验收',                  '2026-06-11','2026-06-13','3天','项目经理'),
    ]
    PHASE_BG = {'1':'E2EFD9','2':'DEEAF1','3':'FFF2CC','4':'FFDEDE','5':'E8D5F5','6':'F2F2F2'}
    for row_data in wbs_rows:
        row = wbs_tbl.add_row()
        pid = row_data[0].split('.')[0]
        bg = PHASE_BG.get(pid, 'FFFFFF')
        is_parent = '.' not in row_data[0]
        for c_idx, val in enumerate(row_data):
            normal_cell(row.cells[c_idx], val, font_size=10,
                        align=WD_ALIGN_PARAGRAPH.CENTER, bold=is_parent)
            if is_parent:
                set_cell_bg(row.cells[c_idx], bg)

    add_heading(doc, '2.2 软件生命周期模型', 2)

    add_heading(doc, '2.2.1 模型图示', 3)
    add_para(doc,
        '本项目采用瀑布模型（Waterfall Model）作为软件生命周期模型。瀑布模型将软件开发过程划分为'
        '若干相互衔接的顺序阶段，各阶段有明确的输入输出文档，适合需求相对稳定的政府信息化项目。'
        '项目流程如下图所示：')
    add_image(doc, images['lifecycle'], width_inches=6.0,
              caption='图2-2  软件生命周期模型（瀑布模型）')

    add_heading(doc, '2.2.2 详细文档（各阶段描述）', 3)
    phases_desc = [
        ('软件规划',     '明确项目目标、可行性分析、项目范围定义、资源与进度初步规划，输出《项目可行性研究报告》和本《项目计划书》。'),
        ('需求开发',     '通过访谈、原型等方式开展业务流程梳理和用例建模，输出《需求规格说明书》，经需求评审后进入下一阶段。'),
        ('软件结构设计', '根据需求设计系统整体架构，包括Agent分层架构设计（EnterpriseAgent、CityAgent、ProvinceAgent、SystemAgent）和REST API接口设计，输出《软件设计说明书》。'),
        ('数据库设计',   '设计系统数据模型，包括ER图、数据字典、表结构定义，输出《数据库设计说明书》，确保数据安全性和完整性。'),
        ('实施（编码）', '按照设计文档实现各Agent模块和REST API层，遵循编码规范，进行代码审查，输出可运行的软件系统源代码。'),
        ('系统集成',     '将各模块集成为完整系统，开展单元测试和集成测试，修复缺陷，输出《测试报告》，确保系统功能符合需求规格。'),
        ('提交上线',     '完成系统部署、配置和用户培训，开展系统验收测试，客户签字确认，输出《系统验收报告》，正式上线运行。'),
        ('维护',         '系统上线后持续跟踪运行状态，处理用户反馈，进行缺陷修复和功能优化，保障系统稳定可靠运行。'),
    ]
    phase_tbl = doc.add_table(rows=1, cols=3)
    phase_tbl.style = 'Table Grid'
    phase_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['序号', '阶段名称', '阶段描述']):
        bold_cell(phase_tbl.rows[0].cells[i], h, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')
    for idx, (phase, desc) in enumerate(phases_desc, 1):
        row = phase_tbl.add_row()
        normal_cell(row.cells[0], str(idx), font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        normal_cell(row.cells[1], phase, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        normal_cell(row.cells[2], desc, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        if idx % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F5F5F5')

    doc.add_page_break()

    # ── Chapter 3: 进度计划 ───────────────────────────────────────────────────
    add_heading(doc, '第三章  进度计划', 1)
    add_para(doc,
        '本章通过甘特图、网络图和里程碑图三种可视化工具展示项目的进度安排。项目总工期为53个日历日，'
        '计划于2026年4月22日开工，2026年6月13日完成客户验收。')

    add_heading(doc, '3.1 甘特图', 2)
    add_para(doc,
        '甘特图以时间轴为横轴，任务列表为纵轴，直观展示各任务的时间安排和并行关系。'
        '图中菱形标记（◆）表示里程碑节点。')
    add_image(doc, images['gantt'], width_inches=6.5,
              caption='图3-1  项目甘特图')

    add_heading(doc, '3.2 网络图', 2)
    add_para(doc,
        '网络图（CPM网络图）展示任务间的逻辑依赖关系，红色箭头为关键路径，沿关键路径任何任务的延误'
        '都将直接影响项目总工期。本项目关键路径总工期为53天。')
    add_image(doc, images['network'], width_inches=6.5,
              caption='图3-2  项目网络图（关键路径）')

    add_heading(doc, '3.3 里程碑图', 2)
    add_para(doc, '项目共设置6个里程碑节点，如下图所示。各里程碑的详细说明见下表。')
    add_image(doc, images['milestone'], width_inches=6.2,
              caption='图3-3  项目里程碑图')

    ms_tbl = doc.add_table(rows=1, cols=4)
    ms_tbl.style = 'Table Grid'
    ms_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['里程碑编号', '里程碑描述', '目标日期', '验收标准']):
        bold_cell(ms_tbl.rows[0].cells[i], h, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')
    ms_data = [
        ('M1', '需求评审通过',       '2026-04-28', '《需求规格说明书》通过评审，获得甲方签字确认'),
        ('M2', '设计文档评审通过',    '2026-05-06', '《软件设计说明书》和《数据库设计说明书》通过评审'),
        ('M3', '编码完成/单元测试通过','2026-05-27', '所有模块代码提交，单元测试覆盖率≥80%，无严重缺陷'),
        ('M4', '集成测试完成',        '2026-06-06', '全部集成测试用例通过，遗留缺陷均为低优先级'),
        ('M5', '系统部署上线',        '2026-06-10', '系统成功部署到生产环境，可正常访问和使用'),
        ('M6', '客户验收通过',        '2026-06-13', '客户完成验收测试，签署《系统验收报告》'),
    ]
    for i, (mid, mdesc, mdate, mcriteria) in enumerate(ms_data):
        row = ms_tbl.add_row()
        for c_idx, val in enumerate([mid, mdesc, mdate, mcriteria]):
            normal_cell(row.cells[c_idx], val, font_size=10,
                        align=WD_ALIGN_PARAGRAPH.CENTER if c_idx < 3 else WD_ALIGN_PARAGRAPH.LEFT)
        if i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F0F5FF')

    doc.add_page_break()

    # ── Chapter 4: 成本计划 ───────────────────────────────────────────────────
    add_heading(doc, '第四章  成本计划', 1)

    add_heading(doc, '4.1 成本估算', 2)
    add_para(doc,
        '本项目采用工作量估算法（基于WBS）进行成本估算。根据各阶段工期和人员配置，估算各类成本如下：')

    cost_tbl = doc.add_table(rows=1, cols=5)
    cost_tbl.style = 'Table Grid'
    cost_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['成本类别', '项目', '工期/数量', '单价（元）', '小计（元）']):
        bold_cell(cost_tbl.rows[0].cells[i], h, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')
    cost_items = [
        ('人力成本', '项目经理（1人×53天）',    '53天', '800', '42,400'),
        ('人力成本', '需求工程师（1人×7天）',    '7天',  '700', '4,900'),
        ('人力成本', '架构师（1人×8天）',        '8天',  '900', '7,200'),
        ('人力成本', '开发工程师A（21天）',       '21天', '750', '15,750'),
        ('人力成本', '开发工程师B（21天）',       '21天', '750', '15,750'),
        ('人力成本', '测试工程师（1人×10天）',   '10天', '700', '7,000'),
        ('人力成本', '运维工程师（1人×4天）',    '4天',  '700', '2,800'),
        ('设备成本', '服务器采购/租用',           '1套',  '—',   '15,000'),
        ('设备成本', '测试设备',                  '1批',  '—',   '3,000'),
        ('软件成本', '开发工具及许可证',          '1套',  '—',   '5,000'),
        ('管理成本', '项目管理及行政支出',        '—',    '—',   '8,000'),
        ('风险储备', '不可预见费（5%）',          '—',    '—',   '6,290'),
        ('合计',    '—',                          '—',    '—',   '133,090'),
    ]
    for i, row_data in enumerate(cost_items):
        row = cost_tbl.add_row()
        is_total = row_data[0] == '合计'
        for c_idx, val in enumerate(row_data):
            normal_cell(row.cells[c_idx], val, font_size=10,
                        align=WD_ALIGN_PARAGRAPH.CENTER, bold=is_total)
        if is_total:
            for c in row.cells:
                set_cell_bg(c, 'FFF2CC')
        elif i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F5F5F5')

    add_heading(doc, '4.2 成本预算', 2)
    add_para(doc,
        '依据成本估算结果，按照项目阶段进行成本预算分配，确保各阶段资金合理使用。成本预算分配方案如下表：')

    budget_tbl = doc.add_table(rows=1, cols=4)
    budget_tbl.style = 'Table Grid'
    budget_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['项目阶段', '计划开始', '计划结束', '预算金额（元）']):
        bold_cell(budget_tbl.rows[0].cells[i], h, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')
    budget_data = [
        ('需求分析阶段',  '2026-04-22', '2026-04-28', '12,000'),
        ('系统设计阶段',  '2026-04-29', '2026-05-06', '15,000'),
        ('编码实现阶段',  '2026-05-07', '2026-05-27', '55,000'),
        ('测试阶段',      '2026-05-28', '2026-06-06', '18,000'),
        ('部署上线阶段',  '2026-06-07', '2026-06-10', '12,000'),
        ('验收阶段',      '2026-06-11', '2026-06-13', '8,000'),
        ('风险储备',      '全程',        '全程',        '6,290'),
        ('设备及软件',    '全程',        '全程',        '23,000'),
        ('合  计',        '—',           '—',           '149,290'),
    ]
    for i, row_data in enumerate(budget_data):
        row = budget_tbl.add_row()
        is_total = '合' in row_data[0]
        for c_idx, val in enumerate(row_data):
            normal_cell(row.cells[c_idx], val, font_size=10,
                        align=WD_ALIGN_PARAGRAPH.CENTER, bold=is_total)
        if is_total:
            for c in row.cells:
                set_cell_bg(c, 'FFF2CC')
        elif i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F5F5F5')

    add_para(doc,
        '注：成本预算包含5%风险储备金和设备软件采购费用，实际执行时按月度执行进度进行成本控制。',
        size=9)

    doc.add_page_break()

    # ── Chapter 5: 人力资源计划 ───────────────────────────────────────────────
    add_heading(doc, '第五章  人力资源计划', 1)

    add_heading(doc, '5.1 项目组织结构', 2)
    add_para(doc,
        '本项目采用职能型组织结构，设置项目经理统筹协调各功能团队。项目组织结构如下图所示：')
    add_image(doc, images['orgchart'], width_inches=6.0,
              caption='图5-1  项目组织结构图')

    add_para(doc, '项目团队成员及职责如下表所示：')
    team_tbl = doc.add_table(rows=1, cols=4)
    team_tbl.style = 'Table Grid'
    team_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['角色', '人数', '参与阶段', '主要职责']):
        bold_cell(team_tbl.rows[0].cells[i], h, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')
    team_data = [
        ('项目经理',    '1', '全程',     '整体计划制定、进度控制、风险管理、与甲方沟通'),
        ('需求工程师',  '1', '需求阶段', '业务调研、流程梳理、需求文档编写、用例建模'),
        ('架构师',      '1', '设计阶段', 'Agent架构设计、数据库设计、技术方案评审'),
        ('开发工程师A', '1', '编码阶段', 'SystemAgent、ProvinceAgent、REST API层开发'),
        ('开发工程师B', '1', '编码阶段', 'EnterpriseAgent、CityAgent开发'),
        ('测试工程师',  '1', '测试阶段', '测试方案制定、单元测试、集成测试执行与报告'),
        ('运维工程师',  '1', '部署阶段', '生产环境搭建、系统部署配置、上线保障'),
    ]
    for i, row_data in enumerate(team_data):
        row = team_tbl.add_row()
        for c_idx, val in enumerate(row_data):
            normal_cell(row.cells[c_idx], val, font_size=10,
                        align=WD_ALIGN_PARAGRAPH.CENTER if c_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT)
        if i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F0F5FF')

    add_heading(doc, '5.2 责任分配矩阵', 2)
    add_para(doc,
        '责任分配矩阵（RAM/RACI）明确各任务的责任归属。R=负责执行，A=审批负责，C=提供咨询，I=知会通知。')

    ram_tbl = doc.add_table(rows=1, cols=8)
    ram_tbl.style = 'Table Grid'
    ram_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ram_headers = ['工作任务', '项目经理', '需求工程师', '架构师', '开发A', '开发B', '测试工程师', '运维工程师']
    for i, h in enumerate(ram_headers):
        bold_cell(ram_tbl.rows[0].cells[i], h, font_size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')
    ram_data = [
        ('需求分析',        'A', 'R', 'C', 'I', 'I', 'I', 'I'),
        ('系统架构设计',    'A', 'C', 'R', 'I', 'I', 'I', 'I'),
        ('数据库设计',      'I', 'I', 'R', 'C', 'C', 'I', 'I'),
        ('SystemAgent开发', 'I', 'I', 'C', 'R', 'I', 'I', 'I'),
        ('Enterprise开发',  'I', 'I', 'C', 'I', 'R', 'I', 'I'),
        ('CityAgent开发',   'I', 'I', 'C', 'I', 'R', 'I', 'I'),
        ('ProvinceAgent开发','I','I', 'C', 'R', 'I', 'I', 'I'),
        ('REST API开发',    'I', 'I', 'C', 'R', 'C', 'I', 'I'),
        ('单元测试',        'I', 'I', 'C', 'C', 'C', 'R', 'I'),
        ('集成测试',        'A', 'I', 'C', 'C', 'C', 'R', 'I'),
        ('部署上线',        'A', 'I', 'C', 'C', 'C', 'C', 'R'),
        ('验收',            'R', 'I', 'I', 'I', 'I', 'C', 'C'),
    ]
    for i, row_data in enumerate(ram_data):
        row = ram_tbl.add_row()
        for c_idx, val in enumerate(row_data):
            bold = (val == 'R')
            normal_cell(row.cells[c_idx], val, font_size=9,
                        align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold)
            if val == 'R':
                set_cell_bg(row.cells[c_idx], 'E2EFD9')
            elif val == 'A':
                set_cell_bg(row.cells[c_idx], 'FFF2CC')
        if i % 2 == 0:
            normal_cell(row.cells[0], row_data[0], font_size=9,
                        align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, '5.3 人力资源管理计划', 2)
    add_para(doc,
        '根据项目各阶段的工作内容，制定人力资源投入计划。下图为项目人力资源直方图，展示各时间段的团队成员投入情况：')
    add_image(doc, images['resource_histogram'], width_inches=6.2,
              caption='图5-2  项目人力资源直方图')

    add_para(doc, '人力资源管理原则：')
    for item in [
        '（1）团队成员招募：项目启动前2周完成团队组建，确保核心成员到岗。',
        '（2）培训计划：项目启动时对团队进行业务背景培训和技术方案培训，统一认知。',
        '（3）绩效考核：按月对团队成员进行绩效评估，考核指标包括任务完成率、质量达标率和协作能力。',
        '（4）人员流动应对：对关键技术文档进行归档管理，实施结对编程，降低人员离职对项目的影响。',
        '（5）团队激励：设置阶段性奖励，里程碑达成后给予团队额外激励，提高团队士气。',
    ]:
        add_para(doc, item, indent=0.5)

    doc.add_page_break()

    # ── Chapter 6: 沟通计划 ───────────────────────────────────────────────────
    add_heading(doc, '第六章  沟通计划', 1)
    add_para(doc,
        '有效的沟通是项目成功的关键因素。本章明确项目各干系人之间的沟通需求、沟通内容、沟通方式和时间安排，'
        '确保信息及时、准确地传达到相关人员。')

    add_heading(doc, '6.1 沟通需求', 2)
    add_para(doc, '项目干系人沟通需求分析如下：')

    comm_need_tbl = doc.add_table(rows=1, cols=4)
    comm_need_tbl.style = 'Table Grid'
    comm_need_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['干系人', '角色', '沟通需求', '关注重点']):
        bold_cell(comm_need_tbl.rows[0].cells[i], h, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')
    need_data = [
        ('云南省劳动和\n社会保障局', '甲方/客户',   '项目进展、里程碑达成情况、变更审批', '进度、质量、费用'),
        ('项目经理',    '项目负责人', '团队工作状态、风险预警、资源协调',         '全局进度与风险'),
        ('开发团队',    '执行者',     '任务分配、技术方案确认、问题澄清',         '技术方案、任务边界'),
        ('测试工程师',  '质量保障',   '测试计划确认、缺陷跟踪、测试报告审批',    '质量目标、缺陷状态'),
        ('运维工程师',  '部署支持',   '部署方案确认、上线时间节点',              '部署方案、环境需求'),
    ]
    for i, row_data in enumerate(need_data):
        row = comm_need_tbl.add_row()
        for c_idx, val in enumerate(row_data):
            normal_cell(row.cells[c_idx], val, font_size=10,
                        align=WD_ALIGN_PARAGRAPH.LEFT)
        if i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F0F5FF')

    add_heading(doc, '6.2 沟通内容', 2)
    add_para(doc, '项目沟通内容分为以下几类：')
    for item in [
        '（1）项目状态报告：每周五提交周报，内容包括本周完成工作、下周计划、存在问题及解决措施。',
        '（2）里程碑汇报：每个里程碑达成后，向甲方提交阶段性工作总结和下阶段计划，并组织评审会议。',
        '（3）变更通知：需求变更、进度调整、资源变动等需及时通知相关干系人，并走正式变更控制流程。',
        '（4）风险预警：识别到高风险事项时，立即向项目经理和甲方汇报，同步应对措施。',
        '（5）技术交流：开发过程中的技术问题讨论、方案评审通过内部会议和即时通讯工具进行。',
    ]:
        add_para(doc, item, indent=0.5)

    add_heading(doc, '6.3 沟通方法', 2)
    add_para(doc, '项目采用多种沟通方式相结合的方式：')

    method_tbl = doc.add_table(rows=1, cols=4)
    method_tbl.style = 'Table Grid'
    method_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['沟通方式', '适用场景', '频率', '负责人']):
        bold_cell(method_tbl.rows[0].cells[i], h, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')
    method_data = [
        ('面对面会议',     '里程碑评审、需求确认、重大决策', '阶段性（约6次）', '项目经理'),
        ('视频/电话会议',  '周例会、问题快速协调',          '每周一次',        '项目经理'),
        ('即时通讯（微信/钉钉）', '日常沟通、快速问答',    '随时',           '全体成员'),
        ('电子邮件',       '正式文档传送、变更通知',        '按需',           '项目经理'),
        ('项目管理系统',   '任务跟踪、进度报告、文档共享',  '每日更新',        '全体成员'),
        ('书面报告',       '周报、月报、阶段总结报告',      '每周/每月',       '项目经理'),
    ]
    for i, row_data in enumerate(method_data):
        row = method_tbl.add_row()
        for c_idx, val in enumerate(row_data):
            normal_cell(row.cells[c_idx], val, font_size=10,
                        align=WD_ALIGN_PARAGRAPH.LEFT)
        if i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F5F5F5')

    add_heading(doc, '6.4 沟通时间安排表', 2)
    add_para(doc, '项目沟通活动时间安排如下表所示：')

    schedule_tbl = doc.add_table(rows=1, cols=5)
    schedule_tbl.style = 'Table Grid'
    schedule_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['沟通活动', '参与人员', '时间/日期', '沟通方式', '产出物']):
        bold_cell(schedule_tbl.rows[0].cells[i], h, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')
    schedule_data = [
        ('项目启动会',         '全体成员+甲方', '2026-04-22', '面对面会议', '会议纪要、项目章程'),
        ('需求评审会',         '甲方+项目经理+需求工程师', '2026-04-28', '面对面会议', '签字确认的需求规格说明书'),
        ('设计评审会',         '项目经理+架构师+甲方', '2026-05-06', '面对面会议', '签字确认的设计文档'),
        ('编码阶段周例会（×3）','全体开发成员', '每周五', '视频会议', '周报'),
        ('测试阶段周例会（×2）','项目经理+测试工程师+开发团队', '每周五', '视频会议', '缺陷跟踪报告'),
        ('部署上线确认会',     '甲方+项目经理+运维', '2026-06-07', '面对面会议', '上线确认单'),
        ('验收会议',           '甲方+全体成员', '2026-06-13', '面对面会议', '系统验收报告'),
        ('项目总结会',         '全体成员',     '2026-06-13', '面对面会议', '项目总结报告'),
    ]
    for i, row_data in enumerate(schedule_data):
        row = schedule_tbl.add_row()
        for c_idx, val in enumerate(row_data):
            normal_cell(row.cells[c_idx], val, font_size=10,
                        align=WD_ALIGN_PARAGRAPH.LEFT)
        if i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F0F5FF')

    add_heading(doc, '6.5 沟通计划维护', 2)
    add_para(doc,
        '沟通计划应根据项目执行情况进行动态调整，具体维护原则如下：')
    for item in [
        '（1）计划审核：每月由项目经理审查沟通计划执行情况，评估沟通效果，必要时调整沟通频率和方式。',
        '（2）变更触发：当项目范围、进度或人员发生重大变化时，应及时更新沟通计划，并通知相关干系人。',
        '（3）文档管理：所有正式沟通记录（会议纪要、评审意见、变更通知等）须归入项目档案库，编号管理。',
        '（4）反馈机制：定期收集干系人对沟通效果的反馈，持续改进沟通质量，确保信息传达的及时性和准确性。',
    ]:
        add_para(doc, item, indent=0.5)

    # ── Appendix: Risk Register ────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, '附录A  风险登记册', 1)
    add_para(doc, '项目识别的主要风险及应对措施如下表所示：')

    risk_tbl = doc.add_table(rows=1, cols=6)
    risk_tbl.style = 'Table Grid'
    risk_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['编号', '风险描述', '可能性', '影响度', '风险级别', '应对措施']):
        bold_cell(risk_tbl.rows[0].cells[i], h, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF', bg='2E75B6')
    risk_data = [
        ('R1', '需求变更频繁',         '高', '中', '中', '建立变更控制流程，需求变更须提交变更申请单，经评审批准后方可实施'),
        ('R2', '开发人员离职',         '中', '高', '高', '关键技术文档化，实施结对编程，建立知识库，降低知识孤岛风险'),
        ('R3', '性能瓶颈（并发填报）', '中', '高', '高', '采用数据库连接池技术，实施压力测试，提前进行性能优化'),
        ('R4', '数据安全泄露',         '低', '极高','高', '数据加密传输存储，角色权限严格隔离，完善日志审计机制'),
        ('R5', '与国家系统对接延迟',   '中', '中', '中', '提前与国家系统主管部门沟通接口规范，设置2周缓冲期'),
    ]
    risk_colors = {'高': 'FFDEDE', '中': 'FFF2CC', '低': 'E2EFD9', '极高': 'FFB3B3'}
    for i, row_data in enumerate(risk_data):
        row = risk_tbl.add_row()
        for c_idx, val in enumerate(row_data):
            normal_cell(row.cells[c_idx], val, font_size=10,
                        align=WD_ALIGN_PARAGRAPH.CENTER if c_idx < 4 else WD_ALIGN_PARAGRAPH.LEFT)
        # Color by risk level
        risk_level = row_data[4]
        bg = risk_colors.get(risk_level, 'FFFFFF')
        for c in row.cells:
            set_cell_bg(c, bg)

    # Save
    out_path = '/home/runner/work/agent/agent/docs/项目计划.docx'
    doc.save(out_path)
    print(f'\n✓ Document saved: {out_path}')
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('Generating project images...')
    images = {
        'gantt':              make_gantt(),
        'network':            make_network(),
        'milestone':          make_milestone(),
        'wbs':                make_wbs(),
        'lifecycle':          make_lifecycle(),
        'orgchart':           make_orgchart(),
        'resource_histogram': make_resource_histogram(),
    }
    print('\nBuilding DOCX document...')
    build_docx(images)
    print('\nAll done!')
