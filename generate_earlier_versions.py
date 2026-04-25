import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.dates as mdates
from matplotlib.patches import FancyArrowPatch
import numpy as np
from datetime import datetime, date, timedelta
import io
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

matplotlib.rcParams['font.sans-serif'] = ['Noto Sans CJK SC', 'Noto Sans CJK', 'SimHei', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False

# ─── DATA DEFINITIONS ─────────────────────────────────────────────────────────

V10_TASKS = [
    # (wbs, name, start, end, duration, owner, is_phase)
    ('1',   '需求分析',        '2026-04-22', '2026-04-28', '7天', '需求工程师', True),
    ('1.1', '业务流程梳理',    '2026-04-22', '2026-04-24', '3天', '需求工程师', False),
    ('1.2', '用例建模',        '2026-04-25', '2026-04-28', '4天', '需求工程师', False),
    ('2',   '系统设计',        '2026-04-29', '2026-05-06', '8天', '架构师', True),
    ('2.1', '架构设计',        '2026-04-29', '2026-05-02', '4天', '架构师', False),
    ('2.2', '数据库设计',      '2026-05-03', '2026-05-06', '4天', '架构师', False),
    ('3',   '编码实现',        '2026-05-07', '2026-05-27', '21天', '开发团队', True),
    ('3.1', 'SystemAgent',     '2026-05-07', '2026-05-10', '4天', '开发A', False),
    ('3.2', 'EnterpriseAgent', '2026-05-11', '2026-05-14', '4天', '开发B', False),
    ('3.3', 'CityAgent',       '2026-05-15', '2026-05-18', '4天', '开发B', False),
    ('3.4', 'ProvinceAgent',   '2026-05-19', '2026-05-23', '5天', '开发A', False),
    ('3.5', 'REST API层',      '2026-05-24', '2026-05-27', '4天', '开发A', False),
    ('4',   '测试',            '2026-05-28', '2026-06-06', '10天', '测试工程师', True),
    ('4.1', '单元测试',        '2026-05-28', '2026-06-01', '5天', '测试工程师', False),
    ('4.2', '集成测试',        '2026-06-02', '2026-06-06', '5天', '测试工程师', False),
    ('5',   '部署上线',        '2026-06-07', '2026-06-10', '4天', '运维', True),
    ('6',   '验收',            '2026-06-11', '2026-06-13', '3天', '项目经理', True),
]

V10_MILESTONES = [
    ('M1', '需求评审通过',           '2026-04-28', '《需求规格说明书》通过评审，获得甲方签字确认'),
    ('M2', '设计文档评审通过',       '2026-05-06', '《软件设计说明书》和《数据库设计说明书》通过评审'),
    ('M3', '编码完成/单元测试通过',  '2026-05-27', '所有模块代码提交，单元测试覆盖率≥80%，无严重缺陷'),
    ('M4', '集成测试完成',           '2026-06-06', '全部集成测试用例通过，遗留缺陷均为低优先级'),
    ('M5', '系统部署上线',           '2026-06-10', '系统成功部署到生产环境，可正常访问和使用'),
    ('M6', '客户验收通过',           '2026-06-13', '客户完成验收测试，签署《系统验收报告》'),
]

V11_TASKS = [
    ('1',     '需求分析',              '2026-04-22', '2026-04-28', '7天',  '需求工程师', True),
    ('1.1',   '业务流程梳理',          '2026-04-22', '2026-04-24', '3天',  '需求工程师', False),
    ('1.2',   '用例建模',              '2026-04-25', '2026-04-28', '4天',  '需求工程师', False),
    ('2',     '系统设计',              '2026-04-29', '2026-05-06', '8天',  '架构师', True),
    ('2.1',   '架构设计',              '2026-04-29', '2026-05-02', '4天',  '架构师', False),
    ('2.2',   '数据库设计',            '2026-05-03', '2026-05-06', '4天',  '架构师', False),
    ('3',     '编码实现',              '2026-05-07', '2026-05-29', '23天', '开发团队', True),
    ('3.1',   'SystemAgent',           '2026-05-07', '2026-05-10', '4天',  '开发A', False),
    ('3.2',   'EnterpriseAgent',       '2026-05-11', '2026-05-14', '4天',  '开发B', False),
    ('3.3',   'CityAgent',             '2026-05-15', '2026-05-18', '4天',  '开发B', False),
    ('3.4',   'ProvinceAgent',         '2026-05-19', '2026-05-23', '5天',  '开发A', False),
    ('3.4.1', '多维分析功能扩展（CR-001）', '2026-05-24', '2026-05-25', '2天', '开发A', False),
    ('3.5',   'REST API层',            '2026-05-26', '2026-05-29', '4天',  '开发A', False),
    ('4',     '测试',                  '2026-05-30', '2026-06-08', '10天', '测试工程师', True),
    ('4.1',   '单元测试',              '2026-05-30', '2026-06-03', '5天',  '测试工程师', False),
    ('4.2',   '集成测试',              '2026-06-04', '2026-06-08', '5天',  '测试工程师', False),
    ('5',     '部署上线',              '2026-06-09', '2026-06-12', '4天',  '运维', True),
    ('6',     '验收',                  '2026-06-13', '2026-06-15', '3天',  '项目经理', True),
]

V11_MILESTONES = [
    ('M1', '需求评审通过',           '2026-04-28', '《需求规格说明书》通过评审，获得甲方签字确认'),
    ('M2', '设计文档评审通过',       '2026-05-06', '《软件设计说明书》和《数据库设计说明书》通过评审'),
    ('M3', '编码完成/单元测试通过',  '2026-05-29', '所有模块代码提交，单元测试覆盖率≥80%，无严重缺陷'),
    ('M4', '集成测试完成',           '2026-06-08', '全部集成测试用例通过，遗留缺陷均为低优先级'),
    ('M5', '系统部署上线',           '2026-06-12', '系统成功部署到生产环境，可正常访问和使用'),
    ('M6', '客户验收通过',           '2026-06-15', '客户完成验收测试，签署《系统验收报告》'),
]

# ─── GANTT CHART GENERATOR ────────────────────────────────────────────────────

def make_gantt(tasks, milestones, title, save_path):
    """Generate Gantt chart and save to save_path."""
    # Filter subtasks for display (skip top-level phases from bar chart)
    display_tasks = [t for t in tasks if not t[6]]  # non-phase tasks
    # Also include phase rows as summary bars
    all_display = tasks[:]

    fig, ax = plt.subplots(figsize=(14, len(all_display) * 0.45 + 2))

    colors = {
        True:  '#2E86AB',   # phase rows
        False: '#A8DADC',   # subtask rows
    }
    cr_color = '#E8A838'    # CR tasks

    y_labels = []
    y_positions = []

    for idx, task in enumerate(all_display):
        wbs, name, start_s, end_s, dur, owner, is_phase = task
        start = datetime.strptime(start_s, '%Y-%m-%d')
        end   = datetime.strptime(end_s,   '%Y-%m-%d') + timedelta(days=1)
        y = len(all_display) - idx - 1

        is_cr = 'CR-' in name
        color = cr_color if is_cr else (colors[True] if is_phase else colors[False])
        height = 0.7 if is_phase else 0.5
        alpha  = 0.95 if is_phase else 0.8

        ax.barh(y, (end - start).days, left=mdates.date2num(start),
                height=height, color=color, alpha=alpha,
                edgecolor='#333', linewidth=0.5)

        label = f'{wbs} {name}'
        y_labels.append(label)
        y_positions.append(y)

    # Milestones
    for ms in milestones:
        mid, mdesc, mdate_s, _ = ms
        mdate = datetime.strptime(mdate_s, '%Y-%m-%d')
        ax.plot(mdates.date2num(mdate), -0.5, 'D', color='red', markersize=8, zorder=5)
        ax.annotate(f'{mid}', (mdates.date2num(mdate), -0.5),
                    textcoords='offset points', xytext=(0, 6),
                    ha='center', fontsize=7, color='red', fontweight='bold')

    ax.set_yticks(y_positions)
    ax.set_yticklabels(y_labels, fontsize=8)
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%m-%d'))
    ax.xaxis.set_major_locator(mdates.WeekdayLocator(byweekday=0))
    ax.xaxis.set_minor_locator(mdates.DayLocator())
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    ax.set_ylim(-1.5, len(all_display))

    # Parse date range
    all_starts = [datetime.strptime(t[2], '%Y-%m-%d') for t in all_display]
    all_ends   = [datetime.strptime(t[3], '%Y-%m-%d') for t in all_display]
    ax.set_xlim(mdates.date2num(min(all_starts) - timedelta(days=2)),
                mdates.date2num(max(all_ends)   + timedelta(days=3)))

    ax.set_title(title, fontsize=12, fontweight='bold', pad=10)
    ax.set_xlabel('日期', fontsize=9)

    legend_patches = [
        mpatches.Patch(color='#2E86AB', label='阶段任务'),
        mpatches.Patch(color='#A8DADC', label='子任务'),
        mpatches.Patch(color='#E8A838', label='CR新增任务'),
        plt.Line2D([0], [0], marker='D', color='w', markerfacecolor='red',
                   markersize=8, label='里程碑'),
    ]
    ax.legend(handles=legend_patches, loc='lower right', fontsize=8)

    plt.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'  Gantt saved to {save_path}')


# ─── NETWORK DIAGRAM GENERATOR ───────────────────────────────────────────────

def make_network(tasks, milestones, title, save_path):
    """Simple CPM-style network diagram."""
    phases = [t for t in tasks if t[6]]
    fig, ax = plt.subplots(figsize=(14, 5))
    ax.set_xlim(-0.5, len(phases) + 0.5)
    ax.set_ylim(-1, 2)
    ax.axis('off')
    ax.set_title(title, fontsize=11, fontweight='bold')

    positions = {}
    for i, t in enumerate(phases):
        x = i + 0.5
        y = 0.5
        positions[t[0]] = (x, y)
        rect = plt.Rectangle((x - 0.4, y - 0.3), 0.8, 0.6, color='#2E86AB',
                               alpha=0.85, ec='#333', lw=1.5)
        ax.add_patch(rect)
        ax.text(x, y + 0.05, t[1], ha='center', va='center', fontsize=8,
                fontweight='bold', color='white', wrap=True)
        ax.text(x, y - 0.18, f"{t[2]}~{t[3]}", ha='center', va='center',
                fontsize=6, color='white')

    # Arrows between sequential phases
    for i in range(len(phases) - 1):
        x1, y1 = positions[phases[i][0]]
        x2, y2 = positions[phases[i+1][0]]
        ax.annotate('', xy=(x2 - 0.42, y2), xytext=(x1 + 0.42, y1),
                    arrowprops=dict(arrowstyle='->', color='red', lw=2))

    # Milestones below
    ms_x = [i + 0.5 for i in range(len(milestones))]
    for i, ms in enumerate(milestones):
        x = ms_x[i] * (len(phases) / len(milestones))
        ax.plot(x, -0.4, 'D', color='red', markersize=10)
        ax.text(x, -0.65, ms[0], ha='center', fontsize=8,
                fontweight='bold', color='red')
        ax.text(x, -0.82, ms[2], ha='center', fontsize=7, color='#555')

    ax.text(0.5, 1.6, '◀─── 关键路径 ───▶', ha='center', fontsize=9,
            color='red', transform=ax.transAxes)

    plt.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'  Network diagram saved to {save_path}')


# ─── MILESTONE CHART GENERATOR ────────────────────────────────────────────────

def make_milestone_chart(milestones, title, save_path):
    fig, ax = plt.subplots(figsize=(12, 4))
    ax.set_xlim(-1, len(milestones))
    ax.set_ylim(-0.5, 1.5)
    ax.axis('off')
    ax.set_title(title, fontsize=11, fontweight='bold')

    # Timeline
    ax.axhline(0.5, color='#555', lw=2, xmin=0.02, xmax=0.98)

    for i, ms in enumerate(milestones):
        mid, desc, dt_s, _ = ms
        x = i
        # Diamond marker
        diamond = plt.Polygon([[x, 0.65], [x+0.12, 0.5], [x, 0.35], [x-0.12, 0.5]],
                               color='red', zorder=5)
        ax.add_patch(diamond)
        # Label alternating above/below
        if i % 2 == 0:
            ax.text(x, 0.75, mid, ha='center', fontsize=9, fontweight='bold', color='red')
            ax.text(x, 0.90, desc, ha='center', fontsize=7.5, color='#333')
            ax.text(x, 1.05, dt_s, ha='center', fontsize=7, color='#666')
        else:
            ax.text(x, 0.22, mid, ha='center', fontsize=9, fontweight='bold', color='red')
            ax.text(x, 0.08, desc, ha='center', fontsize=7.5, color='#333')
            ax.text(x, -0.08, dt_s, ha='center', fontsize=7, color='#666')

    plt.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'  Milestone chart saved to {save_path}')


# ─── ORG CHART GENERATOR ─────────────────────────────────────────────────────

def make_org_chart(save_path):
    fig, ax = plt.subplots(figsize=(12, 5))
    ax.set_xlim(-1, 7)
    ax.set_ylim(-1, 3)
    ax.axis('off')
    ax.set_title('项目组织结构图', fontsize=11, fontweight='bold')

    def box(x, y, text, color='#2E86AB'):
        rect = plt.Rectangle((x-0.55, y-0.22), 1.1, 0.44, color=color, alpha=0.85, ec='#333', lw=1.5)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, color='white', fontweight='bold')

    # PM at top
    box(3, 2.3, '项目经理', '#1A5276')
    # Level 2
    level2 = [(1, 1.2, '需求工程师'), (2.5, 1.2, '架构师'),
              (4, 1.2, '开发A\n开发B'), (5.5, 1.2, '测试工程师')]
    for x, y, name in level2:
        box(x, y, name)
        ax.plot([3, x], [2.08, 1.42], 'k-', lw=1.2, alpha=0.6)
    box(6.5, 1.2, '运维工程师')
    ax.plot([3, 6.5], [2.08, 1.42], 'k-', lw=1.2, alpha=0.6)

    plt.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'  Org chart saved to {save_path}')


# ─── HR HISTOGRAM GENERATOR ──────────────────────────────────────────────────

def make_hr_histogram(tasks, milestones, title, save_path):
    """Monthly HR histogram."""
    # Compute person-days per month
    months_effort = {}
    roles_by_phase = {
        '需求分析': '需求工程师', '系统设计': '架构师',
        '编码实现': '开发团队', '测试': '测试工程师',
        '部署上线': '运维', '验收': '项目经理',
    }
    phases = [t for t in tasks if t[6]]
    data = {}
    for t in phases:
        s = datetime.strptime(t[2], '%Y-%m-%d')
        e = datetime.strptime(t[3], '%Y-%m-%d')
        role = roles_by_phase.get(t[1], t[5])
        days = (e - s).days + 1
        m = f"{s.month}月"
        data.setdefault(m, {})
        data[m][role] = data[m].get(role, 0) + days

    months = sorted(data.keys(), key=lambda x: int(x[:-1]))
    roles = ['需求工程师', '架构师', '开发团队', '测试工程师', '运维', '项目经理']
    role_colors = ['#3498DB', '#2ECC71', '#E74C3C', '#F39C12', '#9B59B6', '#1ABC9C']

    fig, ax = plt.subplots(figsize=(10, 4))
    x = np.arange(len(months))
    width = 0.12
    for i, (role, color) in enumerate(zip(roles, role_colors)):
        vals = [data.get(m, {}).get(role, 0) for m in months]
        ax.bar(x + i * width - 3 * width, vals, width, label=role, color=color, alpha=0.85)

    ax.set_xticks(x)
    ax.set_xticklabels(months)
    ax.set_ylabel('人天数')
    ax.set_title(title, fontsize=11, fontweight='bold')
    ax.legend(fontsize=8, loc='upper right')
    ax.grid(axis='y', alpha=0.3)
    plt.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'  HR histogram saved to {save_path}')


# ─── DOCUMENT BUILDER ────────────────────────────────────────────────────────

def set_cell_shading(cell, fill_hex):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_paragraph(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_image_placeholder(doc, caption, image_path, width_inches=6.0):
    """Insert image and caption paragraph."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cp.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return cp


def style_header_row(table, row_idx=0, fill='2E5F8A'):
    row = table.rows[row_idx]
    for cell in row.cells:
        set_cell_shading(cell, fill)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def style_phase_row(table, row_idx, fill='D6E4F0'):
    row = table.rows[row_idx]
    for cell in row.cells:
        set_cell_shading(cell, fill)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)


def build_doc(version, date_str, tasks, milestones, gantt_path, network_path,
              ms_chart_path, org_path, hr_path, end_date_str, total_days,
              cost_rows, budget_rows, version_note=''):
    """Build complete project plan DOCX."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(11)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run('云南省企业就业失业数据采集系统')
    run.bold = True
    run.font.size = Pt(18)

    sub_title = doc.add_paragraph()
    sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_title.add_run('软件项目计划书')
    run2.bold = True
    run2.font.size = Pt(16)

    doc.add_paragraph()

    cover_table = doc.add_table(rows=7, cols=4)
    cover_table.style = 'Table Grid'
    plan_end = end_date_str
    data = [
        ('项目名称', '云南省企业就业失业数据采集系统', '项目编号', 'YNEM-2026-001'),
        ('委托单位', '云南省劳动和社会保障局',         '版本号',   version),
        ('项目经理', '（待定）',                        '密级',     '内部'),
        ('编制日期', f'2026年{date_str}',              '计划周期', f'2026-04-22 至 {plan_end}'),
        ('编制人员', '项目管理团队',                    '批准人',   '（待定）'),
        ('文档状态', '正式发布',                        '页数',     '共若干页'),
        ('批准签字：_______________          ',
         '批准签字：_______________          ',
         '批准签字：_______________          ',
         '批准签字：_______________          '),
    ]
    for r_idx, row_data in enumerate(data):
        row = cover_table.rows[r_idx]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if c_idx in (0, 2):
                        run.bold = True

    doc.add_page_break()

    # ── CHAPTER 1: 前言 ───────────────────────────────────────────────────────
    add_heading(doc, '第一章  前言', 1)

    add_heading(doc, '1.1 项目开发背景', 2)
    doc.add_paragraph(
        '随着云南省经济社会的持续发展，劳动力市场管理日趋复杂，就业与失业数据的实时监测与上报已成为劳动保障工作的重要环节。'
        '目前，省内各企业的就业失业数据依赖人工填报，存在数据滞后、准确性低、管理效率差等问题。'
    )
    doc.add_paragraph(
        '国家人力资源和社会保障部要求各省建立与国家失业监测系统的数据交换机制，云南省亟需建立一套覆盖省、市、企业三级的'
        '就业失业数据采集、审核、汇总与分析系统，以实现数据的规范化管理和高效利用。'
    )
    doc.add_paragraph(
        '基于上述背景，云南省劳动和社会保障局委托开发"云南省企业就业失业数据采集系统"，采用Agent编程范式，构建企业、市、省、系统'
        '四类智能代理协同工作的软件架构，实现数据采集、审核、汇总、上报的全流程自动化。'
    )

    add_heading(doc, '1.2 项目开发目的', 2)
    doc.add_paragraph('本项目的开发目的包括以下几个方面：')
    for item in [
        '（1）实现企业就业人数的动态监测：支持企业用户实时录入、修改和上报就业失业数据。',
        '（2）建立三级数据管理体系：支持企业、市级、省级用户分级管理数据，实现逐级审核与汇总。',
        '（3）实现数据交换与共享：与国家失业监测系统对接，实现数据的自动上传和同步。',
        '（4）提供数据分析与决策支持：通过多维度数据统计分析，为劳动保障部门决策提供依据。',
        '（5）提升管理效率：以信息化手段替代传统人工填报模式，降低人力成本，提高工作效率。',
    ]:
        doc.add_paragraph(item)

    add_heading(doc, '1.3 项目开发意义', 2)
    doc.add_paragraph('本项目的成功实施将对云南省劳动力市场管理产生深远影响：')
    for item in [
        '• 社会意义：推动云南省就业保障信息化建设，提升政府公共服务能力。',
        '• 经济意义：减少人工数据处理成本，提高政策制定效率，助力劳动力资源的优化配置。',
        '• 技术意义：率先在省级劳动保障系统引入Agent编程范式，为全国同类系统提供参考。',
        '• 管理意义：建立标准化的数据采集和审核流程，强化数据治理，提升数据可信度。',
    ]:
        doc.add_paragraph(item)

    if version_note:
        doc.add_paragraph()
        p = doc.add_paragraph(f'【版本说明】{version_note}')
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x20)

    # ── CHAPTER 2: 范围计划 ───────────────────────────────────────────────────
    add_heading(doc, '第二章  范围计划', 1)

    add_heading(doc, '2.1 项目工作分解结构（WBS）', 2)
    doc.add_paragraph(
        '本项目按照软件工程标准进行工作分解，将项目总体工作分解为若干一级任务包和二级任务包。'
        '下表为项目工作分解结构（WBS）详细说明。'
    )

    # WBS table
    wbs_table = doc.add_table(rows=len(tasks)+1, cols=6)
    wbs_table.style = 'Table Grid'
    headers = ['WBS编号', '任务名称', '计划开始', '计划结束', '工期', '负责人']
    for c, h in enumerate(headers):
        cell = wbs_table.rows[0].cells[c]
        cell.text = h
    style_header_row(wbs_table)

    for r_idx, task in enumerate(tasks, 1):
        wbs, name, start, end, dur, owner, is_phase = task
        row = wbs_table.rows[r_idx]
        vals = [wbs, name, start, end, dur, owner]
        for c, val in enumerate(vals):
            row.cells[c].text = val
            for para in row.cells[c].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        if is_phase:
            style_phase_row(wbs_table, r_idx)
        # CR tasks in amber
        if 'CR-' in name:
            for cell in row.cells:
                set_cell_shading(cell, 'FFF3CD')

    doc.add_paragraph('注：蓝色行为一级阶段任务，白色行为子任务。')
    if version_note and 'CR-001' in version_note:
        doc.add_paragraph('橙黄色行（3.4.1）为CR-001变更新增任务。')

    add_heading(doc, '2.2 软件生命周期模型', 2)
    add_heading(doc, '2.2.1 模型图示', 3)
    doc.add_paragraph(
        '本项目采用瀑布模型（Waterfall Model）作为软件生命周期模型。'
        '瀑布模型将软件开发过程划分为若干相互衔接的顺序阶段，各阶段有明确的输入输出文档，'
        '适合需求相对稳定、系统边界清晰的政务类系统开发。'
    )

    add_heading(doc, '2.2.2 详细文档（各阶段描述）', 3)
    lc_table = doc.add_table(rows=9, cols=3)
    lc_table.style = 'Table Grid'
    lc_headers = ['序号', '阶段名称', '阶段描述']
    for c, h in enumerate(lc_headers):
        lc_table.rows[0].cells[c].text = h
    style_header_row(lc_table)

    lc_data = [
        ('1', '软件规划',    '明确项目目标、可行性分析、项目范围定义、资源与进度初步规划，输出《项目计划书》。'),
        ('2', '需求开发',    '通过访谈、原型等方式开展业务流程梳理和用例建模，输出《需求规格说明书》。'),
        ('3', '软件结构设计','根据需求设计系统整体架构，包括Agent分层架构设计，输出《软件设计说明书》。'),
        ('4', '数据库设计',  '设计系统数据模型，包括ER图、数据字典、表结构定义，输出《数据库设计说明书》。'),
        ('5', '实施（编码）','按照设计文档实现各Agent模块和REST API层，遵循编码规范，输出源代码。'),
        ('6', '系统集成',    '将各模块集成为完整系统，开展单元测试和集成测试，修复缺陷，输出测试报告。'),
        ('7', '提交上线',    '完成系统部署、配置和用户培训，开展系统验收测试，客户签字确认。'),
        ('8', '维护',        '系统上线后持续跟踪运行状态，处理用户反馈，进行缺陷修复和功能优化。'),
    ]
    for r_idx, (no, phase, desc) in enumerate(lc_data, 1):
        row = lc_table.rows[r_idx]
        row.cells[0].text = no
        row.cells[1].text = phase
        row.cells[2].text = desc
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    # ── CHAPTER 3: 进度计划 ───────────────────────────────────────────────────
    add_heading(doc, '第三章  进度计划', 1)
    doc.add_paragraph(
        f'本章通过甘特图、网络图和里程碑图三种可视化工具展示项目的进度安排。'
        f'项目总工期为{total_days}个日历日，计划于2026年4月22日开工，{end_date_str}完成客户验收。'
    )

    add_heading(doc, '3.1 甘特图', 2)
    doc.add_paragraph(
        '甘特图以时间轴为横轴，任务列表为纵轴，直观展示各任务的时间安排和并行关系。'
        '图中菱形标记（◆）表示里程碑节点。'
    )
    add_image_placeholder(doc, '图3-1  项目甘特图', gantt_path, width_inches=6.3)

    add_heading(doc, '3.2 网络图', 2)
    doc.add_paragraph(
        '网络图（CPM网络图）展示任务间的逻辑依赖关系，红色箭头为关键路径，'
        f'沿关键路径任何任务的延误都将直接影响项目总工期。本项目关键路径总工期为{total_days}天。'
    )
    add_image_placeholder(doc, '图3-2  项目网络图（关键路径）', network_path, width_inches=6.3)

    add_heading(doc, '3.3 里程碑图', 2)
    doc.add_paragraph(f'项目共设置{len(milestones)}个里程碑节点，如下图所示。各里程碑的详细说明见下表。')
    add_image_placeholder(doc, '图3-3  项目里程碑图', ms_chart_path, width_inches=6.3)

    # Milestone table
    ms_table = doc.add_table(rows=len(milestones)+1, cols=4)
    ms_table.style = 'Table Grid'
    ms_headers = ['里程碑编号', '里程碑描述', '目标日期', '验收标准']
    for c, h in enumerate(ms_headers):
        ms_table.rows[0].cells[c].text = h
    style_header_row(ms_table)
    for r_idx, (mid, desc, dt, criteria) in enumerate(milestones, 1):
        row = ms_table.rows[r_idx]
        row.cells[0].text = mid
        row.cells[1].text = desc
        row.cells[2].text = dt
        row.cells[3].text = criteria
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    # ── CHAPTER 4: 成本计划 ───────────────────────────────────────────────────
    add_heading(doc, '第四章  成本计划', 1)

    add_heading(doc, '4.1 成本估算', 2)
    doc.add_paragraph(
        '本项目采用工作量估算法（基于WBS）进行成本估算。'
        '根据各阶段工期和人员配置，估算各类成本如下：'
    )

    cost_table = doc.add_table(rows=len(cost_rows)+1, cols=5)
    cost_table.style = 'Table Grid'
    for c, h in enumerate(['成本类别', '项目', '工期/数量', '单价（元）', '小计（元）']):
        cost_table.rows[0].cells[c].text = h
    style_header_row(cost_table)
    for r_idx, row_data in enumerate(cost_rows, 1):
        row = cost_table.rows[r_idx]
        for c, val in enumerate(row_data):
            row.cells[c].text = val
            for para in row.cells[c].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        if row_data[0] == '合计':
            style_phase_row(cost_table, r_idx, 'D6E4F0')

    add_heading(doc, '4.2 成本预算', 2)
    doc.add_paragraph(
        '依据成本估算结果，按照项目阶段进行成本预算分配，确保各阶段资金合理使用。成本预算分配方案如下表：'
    )

    budget_table = doc.add_table(rows=len(budget_rows)+1, cols=4)
    budget_table.style = 'Table Grid'
    for c, h in enumerate(['项目阶段', '计划开始', '计划结束', '预算金额（元）']):
        budget_table.rows[0].cells[c].text = h
    style_header_row(budget_table)
    for r_idx, row_data in enumerate(budget_rows, 1):
        row = budget_table.rows[r_idx]
        for c, val in enumerate(row_data):
            row.cells[c].text = val
            for para in row.cells[c].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        if row_data[0] in ('合  计', '合计'):
            style_phase_row(budget_table, r_idx, 'D6E4F0')

    doc.add_paragraph(
        '注：成本预算包含5%风险储备金和设备软件采购费用，实际执行时按月度执行进度进行成本控制。'
    )

    # ── CHAPTER 5: 人力资源计划 ───────────────────────────────────────────────
    add_heading(doc, '第五章  人力资源计划', 1)

    add_heading(doc, '5.1 项目组织结构', 2)
    doc.add_paragraph(
        '本项目采用职能型组织结构，设置项目经理统筹协调各功能团队。项目组织结构如下图所示：'
    )
    add_image_placeholder(doc, '图5-1  项目组织结构图', org_path, width_inches=5.5)
    doc.add_paragraph('项目团队成员及职责如下表所示：')

    team_table = doc.add_table(rows=8, cols=4)
    team_table.style = 'Table Grid'
    for c, h in enumerate(['角色', '人数', '参与阶段', '主要职责']):
        team_table.rows[0].cells[c].text = h
    style_header_row(team_table)
    team_data = [
        ('项目经理',   '1', '全程',   '整体计划制定、进度控制、风险管理、与甲方沟通'),
        ('需求工程师', '1', '需求阶段', '业务调研、流程梳理、需求文档编写、用例建模'),
        ('架构师',     '1', '设计阶段', 'Agent架构设计、数据库设计、技术方案评审'),
        ('开发工程师A','1', '编码阶段', 'SystemAgent、ProvinceAgent、REST API层开发'),
        ('开发工程师B','1', '编码阶段', 'EnterpriseAgent、CityAgent开发'),
        ('测试工程师', '1', '测试阶段', '测试方案制定、单元测试、集成测试执行与报告'),
        ('运维工程师', '1', '部署阶段', '生产环境搭建、系统部署配置、上线保障'),
    ]
    for r_idx, (role, num, phase, duty) in enumerate(team_data, 1):
        row = team_table.rows[r_idx]
        row.cells[0].text = role
        row.cells[1].text = num
        row.cells[2].text = phase
        row.cells[3].text = duty
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    add_heading(doc, '5.2 责任分配矩阵', 2)
    doc.add_paragraph(
        '责任分配矩阵（RAM/RACI）明确各任务的责任归属。R=负责执行，A=审批负责，C=提供咨询，I=知会通知。'
    )

    raci_rows_data = [
        ('需求分析',       'A', 'R', 'C', 'I', 'I', 'I', 'I'),
        ('系统架构设计',   'A', 'C', 'R', 'I', 'I', 'I', 'I'),
        ('数据库设计',     'I', 'I', 'R', 'C', 'C', 'I', 'I'),
        ('SystemAgent开发','I', 'I', 'C', 'R', 'I', 'I', 'I'),
        ('Enterprise开发', 'I', 'I', 'C', 'I', 'R', 'I', 'I'),
        ('CityAgent开发',  'I', 'I', 'C', 'I', 'R', 'I', 'I'),
        ('ProvinceAgent开发','I','I', 'C', 'R', 'I', 'I', 'I'),
        ('REST API开发',   'I', 'I', 'C', 'R', 'C', 'I', 'I'),
        ('单元测试',       'I', 'I', 'C', 'C', 'C', 'R', 'I'),
        ('集成测试',       'A', 'I', 'C', 'C', 'C', 'R', 'I'),
        ('部署上线',       'A', 'I', 'C', 'C', 'C', 'C', 'R'),
        ('验收',           'R', 'I', 'I', 'I', 'I', 'C', 'C'),
    ]
    raci_table = doc.add_table(rows=len(raci_rows_data)+1, cols=8)
    raci_table.style = 'Table Grid'
    for c, h in enumerate(['工作任务', '项目经理', '需求工程师', '架构师',
                            '开发A', '开发B', '测试工程师', '运维工程师']):
        raci_table.rows[0].cells[c].text = h
    style_header_row(raci_table)
    for r_idx, row_data in enumerate(raci_rows_data, 1):
        row = raci_table.rows[r_idx]
        for c, val in enumerate(row_data):
            row.cells[c].text = val
            for para in row.cells[c].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
            if c > 0 and val == 'R':
                set_cell_shading(row.cells[c], 'D5E8D4')
            elif c > 0 and val == 'A':
                set_cell_shading(row.cells[c], 'DAE8FC')

    add_heading(doc, '5.3 人力资源管理计划', 2)
    doc.add_paragraph(
        '根据项目各阶段的工作内容，制定人力资源投入计划。'
        '下图为项目人力资源直方图，展示各时间段的团队成员投入情况：'
    )
    add_image_placeholder(doc, '图5-2  项目人力资源直方图', hr_path, width_inches=5.5)
    doc.add_paragraph('人力资源管理原则：')
    for item in [
        '（1）团队成员招募：项目启动前2周完成团队组建，确保核心成员到岗。',
        '（2）培训计划：项目启动时对团队进行业务背景培训和技术方案培训，统一认知。',
        '（3）绩效考核：按月对团队成员进行绩效评估，考核指标包括任务完成率、质量达标率和协作能力。',
        '（4）人员流动应对：对关键技术文档进行归档管理，实施结对编程，降低人员离职对项目的影响。',
        '（5）团队激励：设置阶段性奖励，里程碑达成后给予团队额外激励，提高团队士气。',
    ]:
        doc.add_paragraph(item)

    # ── CHAPTER 6: 沟通计划 ───────────────────────────────────────────────────
    add_heading(doc, '第六章  沟通计划', 1)
    doc.add_paragraph(
        '有效的沟通是项目成功的关键因素。本章明确项目各干系人之间的沟通需求、沟通内容、'
        '沟通方式和时间安排，确保信息及时、准确地传达到相关人员。'
    )

    add_heading(doc, '6.1 沟通需求', 2)
    doc.add_paragraph('项目干系人沟通需求分析如下：')
    stakeholder_table = doc.add_table(rows=6, cols=4)
    stakeholder_table.style = 'Table Grid'
    for c, h in enumerate(['干系人', '角色', '沟通需求', '关注重点']):
        stakeholder_table.rows[0].cells[c].text = h
    style_header_row(stakeholder_table)
    sh_data = [
        ('云南省劳动和\n社会保障局', '甲方/客户',  '项目进展、里程碑达成情况、变更审批', '进度、质量、费用'),
        ('项目经理',               '项目负责人', '团队工作状态、风险预警、资源协调',   '全局进度与风险'),
        ('开发团队',               '执行者',     '任务分配、技术方案确认、问题澄清',   '技术方案、任务边界'),
        ('测试工程师',             '质量保障',   '测试计划确认、缺陷跟踪、测试报告审批','质量目标、缺陷状态'),
        ('运维工程师',             '部署支持',   '部署方案确认、上线时间节点',         '部署方案、环境需求'),
    ]
    for r_idx, row_data in enumerate(sh_data, 1):
        row = stakeholder_table.rows[r_idx]
        for c, val in enumerate(row_data):
            row.cells[c].text = val
            for para in row.cells[c].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    add_heading(doc, '6.2 沟通内容', 2)
    doc.add_paragraph('项目沟通内容分为以下几类：')
    for item in [
        '（1）项目状态报告：每周五提交周报，内容包括本周完成工作、下周计划、存在问题及解决措施。',
        '（2）里程碑汇报：每个里程碑达成后，向甲方提交阶段性工作总结和下阶段计划，并组织评审会议。',
        '（3）变更通知：需求变更、进度调整、资源变动等需及时通知相关干系人，并走正式变更控制流程。',
        '（4）风险预警：识别到高风险事项时，立即向项目经理和甲方汇报，同步应对措施。',
        '（5）技术交流：开发过程中的技术问题讨论、方案评审通过内部会议和即时通讯工具进行。',
    ]:
        doc.add_paragraph(item)

    add_heading(doc, '6.3 沟通方法', 2)
    doc.add_paragraph('项目采用多种沟通方式相结合的方式：')
    comm_table = doc.add_table(rows=7, cols=4)
    comm_table.style = 'Table Grid'
    for c, h in enumerate(['沟通方式', '适用场景', '频率', '负责人']):
        comm_table.rows[0].cells[c].text = h
    style_header_row(comm_table)
    comm_data = [
        ('面对面会议',           '里程碑评审、需求确认、重大决策', '阶段性（约6次）', '项目经理'),
        ('视频/电话会议',        '周例会、问题快速协调',          '每周一次',       '项目经理'),
        ('即时通讯（微信/钉钉）','日常沟通、快速问答',            '随时',           '全体成员'),
        ('电子邮件',             '正式文档传送、变更通知',        '按需',           '项目经理'),
        ('项目管理系统',         '任务跟踪、进度报告、文档共享',  '每日更新',       '全体成员'),
        ('书面报告',             '周报、月报、阶段总结报告',      '每周/每月',      '项目经理'),
    ]
    for r_idx, row_data in enumerate(comm_data, 1):
        row = comm_table.rows[r_idx]
        for c, val in enumerate(row_data):
            row.cells[c].text = val
            for para in row.cells[c].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    add_heading(doc, '6.4 沟通时间安排表', 2)
    doc.add_paragraph('项目沟通活动时间安排如下表所示：')
    sched_table = doc.add_table(rows=9, cols=5)
    sched_table.style = 'Table Grid'
    for c, h in enumerate(['沟通活动', '参与人员', '时间/日期', '沟通方式', '产出物']):
        sched_table.rows[0].cells[c].text = h
    style_header_row(sched_table)
    phase_end = milestones[-1][2]
    sched_data = [
        ('项目启动会',             '全体成员+甲方',          '2026-04-22',  '面对面会议', '会议纪要、项目章程'),
        ('需求评审会',             '甲方+项目经理+需求工程师','2026-04-28',  '面对面会议', '签字确认的需求规格说明书'),
        ('设计评审会',             '项目经理+架构师+甲方',    '2026-05-06',  '面对面会议', '签字确认的设计文档'),
        ('编码阶段周例会（×3）',   '全体开发成员',           '每周五',      '视频会议',   '周报'),
        ('测试阶段周例会（×2）',   '项目经理+测试工程师+开发','每周五',      '视频会议',   '缺陷跟踪报告'),
        ('部署上线确认会',         '甲方+项目经理+运维',      milestones[4][2], '面对面会议', '上线确认单'),
        ('验收会议',               '甲方+全体成员',          phase_end,     '面对面会议', '系统验收报告'),
        ('项目总结会',             '全体成员',               phase_end,     '面对面会议', '项目总结报告'),
    ]
    for r_idx, row_data in enumerate(sched_data, 1):
        row = sched_table.rows[r_idx]
        for c, val in enumerate(row_data):
            row.cells[c].text = val
            for para in row.cells[c].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    add_heading(doc, '6.5 沟通计划维护', 2)
    doc.add_paragraph(
        '沟通计划应根据项目执行情况进行动态调整，具体维护原则如下：'
    )
    for item in [
        '（1）计划审核：每月由项目经理审查沟通计划执行情况，评估沟通效果，必要时调整频率和方式。',
        '（2）变更触发：当项目范围、进度或人员发生重大变化时，应及时更新沟通计划，并通知相关干系人。',
        '（3）文档管理：所有正式沟通记录（会议纪要、评审意见、变更通知等）须归入项目档案库，编号管理。',
        '（4）反馈机制：定期收集干系人对沟通效果的反馈，持续改进沟通质量，确保信息传达的及时性和准确性。',
    ]:
        doc.add_paragraph(item)

    # ── APPENDIX A: 风险登记册 ────────────────────────────────────────────────
    add_heading(doc, '附录A  风险登记册', 1)
    doc.add_paragraph('项目识别的主要风险及应对措施如下表所示：')

    risk_table = doc.add_table(rows=6, cols=6)
    risk_table.style = 'Table Grid'
    for c, h in enumerate(['编号', '风险描述', '可能性', '影响度', '风险级别', '应对措施']):
        risk_table.rows[0].cells[c].text = h
    style_header_row(risk_table)
    risk_data = [
        ('R1', '需求变更频繁',       '高', '中', '中', '建立变更控制流程，需求变更须提交变更申请单，经评审批准后方可实施。'),
        ('R2', '开发人员离职',       '中', '高', '高', '关键技术文档化，实施结对编程，建立知识库，降低知识孤岛风险。'),
        ('R3', '性能瓶颈（并发填报）','中', '高', '高', '采用数据库连接池技术，实施压力测试，提前进行性能优化。'),
        ('R4', '数据安全泄露',       '低', '极高','高', '数据加密传输存储，角色权限严格隔离，完善日志审计机制。'),
        ('R5', '与国家系统对接延迟', '中', '中', '中', '提前与国家系统主管部门沟通接口规范，设置2周缓冲期。'),
    ]
    for r_idx, row_data in enumerate(risk_data, 1):
        row = risk_table.rows[r_idx]
        for c, val in enumerate(row_data):
            row.cells[c].text = val
            for para in row.cells[c].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        level = row_data[4]
        if level == '高':
            set_cell_shading(row.cells[4], 'FFD7D7')
        elif level == '中':
            set_cell_shading(row.cells[4], 'FFF3CD')

    return doc


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    docs_dir = '/home/runner/work/agent/agent/docs'
    img_dir  = docs_dir   # save images alongside docs

    os.makedirs(img_dir, exist_ok=True)

    # ── V1.0 ──────────────────────────────────────────────────────────────────
    print('Generating V1.0...')
    v10_gantt   = os.path.join(img_dir, '_v10_gantt.png')
    v10_net     = os.path.join(img_dir, '_v10_network.png')
    v10_ms      = os.path.join(img_dir, '_v10_milestone.png')
    v10_org     = os.path.join(img_dir, '_v10_org.png')
    v10_hr      = os.path.join(img_dir, '_v10_hr.png')

    make_gantt(V10_TASKS, V10_MILESTONES, 'V1.0 项目甘特图', v10_gantt)
    make_network(V10_TASKS, V10_MILESTONES, 'V1.0 项目网络图（关键路径）', v10_net)
    make_milestone_chart(V10_MILESTONES, 'V1.0 项目里程碑图', v10_ms)
    make_org_chart(v10_org)
    make_hr_histogram(V10_TASKS, V10_MILESTONES, 'V1.0 人力资源直方图', v10_hr)

    v10_cost_rows = [
        ('人力成本', '项目经理（1人×53天）',    '53天', '800', '42,400'),
        ('人力成本', '需求工程师（1人×7天）',    '7天',  '700', '4,900'),
        ('人力成本', '架构师（1人×8天）',        '8天',  '900', '7,200'),
        ('人力成本', '开发工程师A（21天）',      '21天', '750', '15,750'),
        ('人力成本', '开发工程师B（21天）',      '21天', '750', '15,750'),
        ('人力成本', '测试工程师（1人×10天）',   '10天', '700', '7,000'),
        ('人力成本', '运维工程师（1人×4天）',    '4天',  '700', '2,800'),
        ('设备成本', '服务器采购/租用',          '1套',  '—',  '15,000'),
        ('设备成本', '测试设备',                '1批',  '—',  '3,000'),
        ('软件成本', '开发工具及许可证',         '1套',  '—',  '5,000'),
        ('管理成本', '项目管理及行政支出',       '—',   '—',  '8,000'),
        ('风险储备', '不可预见费（5%）',         '—',   '—',  '6,290'),
        ('合计',     '—',                       '—',   '—',  '133,090'),
    ]
    v10_budget_rows = [
        ('需求分析阶段', '2026-04-22', '2026-04-28', '12,000'),
        ('系统设计阶段', '2026-04-29', '2026-05-06', '15,000'),
        ('编码实现阶段', '2026-05-07', '2026-05-27', '55,000'),
        ('测试阶段',     '2026-05-28', '2026-06-06', '18,000'),
        ('部署上线阶段', '2026-06-07', '2026-06-10', '12,000'),
        ('验收阶段',     '2026-06-11', '2026-06-13', '8,000'),
        ('风险储备',     '全程',       '全程',        '6,290'),
        ('设备及软件',   '全程',       '全程',        '23,000'),
        ('合  计',       '—',         '—',           '149,290'),
    ]

    doc10 = build_doc(
        version='V1.0', date_str='4月22日',
        tasks=V10_TASKS, milestones=V10_MILESTONES,
        gantt_path=v10_gantt, network_path=v10_net,
        ms_chart_path=v10_ms, org_path=v10_org, hr_path=v10_hr,
        end_date_str='2026-06-13', total_days=53,
        cost_rows=v10_cost_rows, budget_rows=v10_budget_rows,
        version_note='',
    )
    out10 = os.path.join(docs_dir, '项目计划_V1.0.docx')
    doc10.save(out10)
    print(f'  Saved: {out10}')

    # ── V1.1 ──────────────────────────────────────────────────────────────────
    print('Generating V1.1...')
    v11_gantt = os.path.join(img_dir, '_v11_gantt.png')
    v11_net   = os.path.join(img_dir, '_v11_network.png')
    v11_ms    = os.path.join(img_dir, '_v11_milestone.png')
    v11_org   = os.path.join(img_dir, '_v11_org.png')
    v11_hr    = os.path.join(img_dir, '_v11_hr.png')

    make_gantt(V11_TASKS, V11_MILESTONES, 'V1.1 项目甘特图（含CR-001）', v11_gantt)
    make_network(V11_TASKS, V11_MILESTONES, 'V1.1 项目网络图（关键路径）', v11_net)
    make_milestone_chart(V11_MILESTONES, 'V1.1 项目里程碑图', v11_ms)
    make_org_chart(v11_org)
    make_hr_histogram(V11_TASKS, V11_MILESTONES, 'V1.1 人力资源直方图', v11_hr)

    v11_cost_rows = [
        ('人力成本', '项目经理（1人×55天）',    '55天', '800', '44,000'),
        ('人力成本', '需求工程师（1人×7天）',    '7天',  '700', '4,900'),
        ('人力成本', '架构师（1人×8天）',        '8天',  '900', '7,200'),
        ('人力成本', '开发工程师A（23天）',      '23天', '750', '17,250'),
        ('人力成本', '开发工程师B（8天）',       '8天',  '750', '6,000'),
        ('人力成本', '测试工程师（1人×10天）',   '10天', '700', '7,000'),
        ('人力成本', '运维工程师（1人×4天）',    '4天',  '700', '2,800'),
        ('设备成本', '服务器采购/租用',          '1套',  '—',  '15,000'),
        ('设备成本', '测试设备',                '1批',  '—',  '3,000'),
        ('软件成本', '开发工具及许可证',         '1套',  '—',  '5,000'),
        ('管理成本', '项目管理及行政支出',       '—',   '—',  '8,000'),
        ('风险储备', '不可预见费（5%）',         '—',   '—',  '6,508'),
        ('合计',     '—',                       '—',   '—',  '136,658'),
    ]
    v11_budget_rows = [
        ('需求分析阶段', '2026-04-22', '2026-04-28', '12,000'),
        ('系统设计阶段', '2026-04-29', '2026-05-06', '15,000'),
        ('编码实现阶段', '2026-05-07', '2026-05-29', '57,500'),
        ('测试阶段',     '2026-05-30', '2026-06-08', '18,000'),
        ('部署上线阶段', '2026-06-09', '2026-06-12', '12,000'),
        ('验收阶段',     '2026-06-13', '2026-06-15', '8,000'),
        ('风险储备',     '全程',       '全程',        '6,508'),
        ('设备及软件',   '全程',       '全程',        '23,000'),
        ('合  计',       '—',         '—',           '152,008'),
    ]

    doc11 = build_doc(
        version='V1.1', date_str='5月3日',
        tasks=V11_TASKS, milestones=V11_MILESTONES,
        gantt_path=v11_gantt, network_path=v11_net,
        ms_chart_path=v11_ms, org_path=v11_org, hr_path=v11_hr,
        end_date_str='2026-06-15', total_days=55,
        cost_rows=v11_cost_rows, budget_rows=v11_budget_rows,
        version_note='本版本依据CR-001（多维分析功能扩展变更单，2026-05-03批准）更新，'
                     '新增WBS任务3.4.1，编码阶段延长2个工作日，里程碑M3~M6相应顺延。',
    )
    out11 = os.path.join(docs_dir, '项目计划_V1.1.docx')
    doc11.save(out11)
    print(f'  Saved: {out11}')

    # Cleanup temp images
    for p in [v10_gantt, v10_net, v10_ms, v10_org, v10_hr,
              v11_gantt, v11_net, v11_ms, v11_org, v11_hr]:
        if os.path.exists(p):
            os.remove(p)
    print('Done. Temp images cleaned up.')


if __name__ == '__main__':
    main()
